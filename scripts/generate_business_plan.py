"""
Generate Impuestify Business Plan as professional Word document.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "plans")

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._tc.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)

def add_styled_table(doc, headers, rows, col_widths=None, highlight_col=None):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1a56db')

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if highlight_col is not None and c_idx == highlight_col:
                        run.bold = True
                        run.font.color.rgb = RGBColor(26, 86, 219)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'f0f5ff')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def build_document():
    doc = Document()

    # === Page margins ===
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Montserrat'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(30, 30, 30)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Montserrat'
        hs.font.color.rgb = RGBColor(26, 86, 219)

    for ls in ['List Bullet', 'List Number']:
        if ls in doc.styles:
            doc.styles[ls].font.name = 'Montserrat'
            doc.styles[ls].font.size = Pt(11)

    # === COVER PAGE ===
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('IMPUESTIFY')
    run.font.size = Pt(42)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Tu Asistente Fiscal Inteligente')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Business Plan — Abril 2026 (v3.0)')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(26, 86, 219)

    doc.add_paragraph()
    doc.add_paragraph()

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run('CONFIDENCIAL')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.italic = True

    doc.add_page_break()

    # === TABLE OF CONTENTS ===
    doc.add_heading('Índice', level=1)
    toc_items = [
        '1. Resumen Ejecutivo',
        '2. El Problema',
        '3. La Solución: Impuestify',
        '4. Producto y Tecnología',
        '5. Mercado Objetivo',
        '6. Análisis Competitivo',
        '7. Modelo de Negocio',
        '8. Tracción y Métricas',
        '9. Roadmap',
        '10. Equipo',
        '11. Financiación y Uso de Fondos',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # === 1. RESUMEN EJECUTIVO ===
    doc.add_heading('1. Resumen Ejecutivo', level=1)

    doc.add_paragraph(
        'Impuestify es el primer asistente fiscal con Inteligencia Artificial multi-agente '
        'que cubre todos los territorios de España, incluyendo los cuatro territorios forales '
        '(País Vasco y Navarra) que ningún competidor atiende.'
    )

    doc.add_paragraph(
        'En España, 21,8 millones de personas presentan la declaración de la renta cada año. '
        'De ellos, 1,6 millones residen en territorios forales y están completamente desatendidos '
        'por las plataformas digitales existentes. Impuestify es la única solución tecnológica '
        'que les ofrece asistencia fiscal avanzada con IA.'
    )

    doc.add_heading('Cifras clave', level=2)
    add_styled_table(doc,
        ['Métrica', 'Valor'],
        [
            ['Deducciones fiscales cubiertas', '~1.008 (16 estatales + 195 territoriales + 339 XSD + 50 forales + 408 CCAA 2025 en 21 territorios)'],
            ['Territorios cubiertos', '21 (17 CCAA + 4 forales + Ceuta/Melilla)'],
            ['Documentos oficiales en RAG', '456+ (PDF + Excel + especificaciones AEAT)'],
            ['Agentes IA especializados', '6 agentes especializados + 1 Coordinator'],
            ['Tests automatizados', '1.702 (100 % passing)'],
            ['Stack tecnológico', 'FastAPI + React + OpenAI GPT + Gemini 3 Flash Vision + RAG avanzado'],
            ['Simulador IRPF completo', 'Trabajo, ahorro, inmuebles, plusvalía, GP, ISD 21 CCAA, familia, deducciones, retenciones, forales, Ceuta/Melilla, 720/721'],
            ['Casillas IRPF Modelo 100', '2.064 casillas oficiales indexadas'],
            ['Clasificador de Facturas IA', 'OCR con Gemini 3 Flash Vision + clasificación PGC automática (~0,0003 $/factura, 33x más barato que Azure DI)'],
            ['Contabilidad PGC', 'Libro Diario, Libro Mayor, Balance, PyG. 66 cuentas PGC (grupos 1-7). Export CSV/Excel para Registro Mercantil'],
            ['Endpoints API', '50+ (incluidos 5 facturas + 5 contabilidad)'],
            ['Calculadoras especializadas', '15+ (IRPF, IVA, RETA, ISD, IPSI, deducciones, Mod.303/130, casillas, perfil fiscal, sueldo neto, retenciones IRPF, conjunta vs. individual, 720/721, múltiples pagadores)'],
            ['Guía Fiscal adaptativa', 'Wizard adaptativo por rol: 7 pasos (particular), 8 pasos (creator con plataformas/IAE/IVA), 8 pasos (autónomo)'],
            ['Calculadora Sueldo Neto', 'ÚNICA en España con 5 regímenes fiscales: IVA, IGIC, IPSI, foral vasco, foral navarro. Cuota SS auto-calculada por ingresos reales (15 tramos)'],
            ['Calculadora Retenciones IRPF', 'Algoritmo oficial AEAT 2026. Gratuita (lead magnet SEO). 28 tests'],
            ['Compartir conversaciones', 'Enlaces públicos con anonimización automática de PII (DNI, IBAN, importes)'],
            ['Sistema de feedback', 'Widget flotante + valoración por respuesta + dashboard admin'],
            ['Plan Particular', '5 €/mes'],
            ['Plan Creator', '49 €/mes (YouTubers, streamers, influencers)'],
            ['Plan Autónomo', '39 €/mes (IVA incluido)'],
        ],
        highlight_col=1
    )

    doc.add_page_break()

    # === 2. EL PROBLEMA ===
    doc.add_heading('2. El Problema', level=1)

    doc.add_heading('2.1 Complejidad fiscal en España', level=2)
    doc.add_paragraph(
        'España tiene uno de los sistemas fiscales más complejos de Europa. Con 17 comunidades '
        'autónomas con competencias fiscales propias, 4 territorios forales con sistemas IRPF '
        'completamente independientes (Araba, Bizkaia, Gipuzkoa y Navarra), y las particularidades '
        'de Ceuta y Melilla, el ciudadano medio se enfrenta a un laberinto normativo.'
    )

    problems = [
        ('El 72% de los españoles', 'no conoce todas las deducciones a las que tiene derecho (INE 2024).'),
        ('1,6 millones de declarantes forales', 'no tienen ninguna herramienta digital avanzada. '
         'TaxDown, líder del mercado, los rechaza explícitamente.'),
        ('3,4 millones de autónomos', 'gestionan trimestrales (IVA, IRPF) con asesores humanos a 50-150 €/mes, '
         'con tiempos de respuesta de 24-48 horas.'),
        ('Las notificaciones de Hacienda', 'generan ansiedad y confusión. No hay herramientas que las analicen '
         'automáticamente con IA.'),
    ]
    for title, desc in problems:
        p = doc.add_paragraph()
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('2.2 Las soluciones actuales son insuficientes', level=2)
    doc.add_paragraph(
        'TaxDown (líder con 1M+ usuarios y 29,7M € recaudados) es un formulario guiado con '
        'revisión humana. Su IA (AsesorIA) es un chatbot básico sobre ChatGPT sin documentos '
        'fiscales reales. No cubre forales, no analiza nóminas automáticamente, y su modelo '
        'depende de 200+ asesores humanos con costes crecientes.'
    )

    doc.add_page_break()

    # === 3. LA SOLUCIÓN ===
    doc.add_heading('3. La Solución: Impuestify', level=1)

    doc.add_paragraph(
        'Impuestify reemplaza al asesor fiscal humano con un sistema multi-agente de IA que:'
    )

    features = [
        'Responde consultas fiscales en segundos (no en 24-48 horas)',
        'Cubre TODOS los territorios de España, incluyendo forales',
        'Analiza nóminas en PDF automáticamente y proyecta IRPF anual',
        'Interpreta notificaciones de Hacienda con extracción estructurada',
        'Descubre deducciones personalizadas entre ~1.000 opciones en 21 territorios',
        'Adjunta documentos directamente en el chat para análisis instantáneo',
        'Calcula el Impuesto de Sucesiones y Donaciones (ISD) en las 21 CCAA',
        'Calcula IPSI trimestral para residentes en Ceuta y Melilla',
        'Calculadora de retenciones IRPF con algoritmo oficial AEAT (gratuita)',
        'Comparador de declaración conjunta vs. individual (4 escenarios)',
        'Compartir conversaciones con enlaces públicos y anonimización PII',
        'Ganancias patrimoniales por inmuebles, plusvalía y renta imputada',
        'Modelos 720/721 para bienes y criptomonedas en el extranjero',
        'Múltiples pagadores: obligación de declarar según Art. 96 LIRPF',
        'Clasificador de facturas con IA: OCR automático (Gemini 3 Flash Vision) + clasificación PGC',
        'Contabilidad PGC completa: Libro Diario, Libro Mayor, Balance de Situación, Cuenta de PyG',
        'Exportación contable CSV/Excel para depositar en el Registro Mercantil',
        '66 cuentas PGC precargadas (grupos 1-7) con asignación automática por concepto',
        'Calendario fiscal personalizado con notificaciones de plazos',
        'Actualiza tu perfil fiscal automáticamente desde el chat',
        'Sistema de feedback integrado (widget + valoración por respuesta)',
        'Calcula IRPF, IVA (Modelo 303), pagos fraccionados (Modelo 130), cuota RETA, ISD e IPSI',
        'Cita fuentes legales en cada respuesta (456+ archivos fiscales oficiales)',
        'Funciona 24/7 sin coste marginal por consulta',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('Propuesta de valor única', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '"El único asistente fiscal con IA que cubre País Vasco y Navarra. '
        'Porque si vives en territorio foral, mereces las mismas ventajas que el resto."'
    )
    run.italic = True
    run.font.color.rgb = RGBColor(26, 86, 219)

    doc.add_page_break()

    # === 4. PRODUCTO Y TECNOLOGIA ===
    doc.add_heading('4. Producto y Tecnología', level=1)

    doc.add_heading('4.1 Arquitectura Multi-Agente', level=2)
    doc.add_paragraph(
        'Impuestify utiliza una arquitectura de agentes especializados que colaboran para resolver '
        'consultas complejas. Cada agente es experto en su dominio y tiene acceso a herramientas específicas.'
    )

    add_styled_table(doc,
        ['Agente', 'Función', 'Herramientas'],
        [
            ['Coordinador', 'Router inteligente, clasifica consultas', 'Routing automático a agente especializado'],
            ['Agente Fiscal', 'Consultas fiscales generales, IRPF, IVA, deducciones, ISD, IPSI',
             'calculate_irpf, discover_deductions, simulate_irpf, calculate_modelo_303, calculate_isd, calculate_modelo_ipsi, lookup_casilla, update_fiscal_profile'],
            ['Agente de Nóminas', 'Análisis de nóminas en PDF',
             'Extracción de 35 campos, proyección anual, detección errores retención'],
            ['Agente de Notificaciones', 'Interpretación notificaciones AEAT',
             'Extracción tipo, plazo, importe, acción requerida'],
            ['Agente Documental', 'Análisis de documentos del usuario',
             'Contexto aislado por workspace, embeddings por documento'],
            ['Agente de Sesión', 'Gestión de documentos efímeros en el chat',
             'Upload, extracción, clasificación automática, anonimización PII'],
            ['Agente Contable', 'Clasificación de facturas y contabilidad PGC',
             'OCR Gemini 3 Flash Vision, clasificación PGC automática, Libro Diario/Mayor, Balance, PyG, export CSV/Excel'],
        ]
    )

    doc.add_heading('4.2 RAG (Retrieval-Augmented Generation)', level=2)
    doc.add_paragraph(
        'Cada respuesta se genera a partir de 456+ archivos fiscales oficiales indexados: '
        'normativa AEAT, BOE, leyes de CCAA, Normas Forales de cada Diputación, y '
        'manuales prácticos. El sistema realiza búsqueda híbrida (semántica + BM25 + FTS5) '
        'para recuperar los fragmentos más relevantes y cita la fuente en cada respuesta.'
    )

    doc.add_heading('4.3 Seguridad y Guardrails IA', level=2)
    security = [
        'Llama Guard 4 — moderación de contenido en 14 categorías',
        'Prompt Guard 2 — detección de inyección de prompts',
        'Filtro PII — detección de DNI, teléfono, email, cuentas bancarias',
        'Anonimización automática de datos personales en documentos adjuntos',
        'Clasificación inteligente de documentos por contenido',
        'Validador SQL — prevención de inyección SQL',
        'Rate limiting — protección contra abuso con cache en memoria',
        'Cache semántico — reducción de costes ~30%',
        'JWT + refresh tokens — autenticación segura',
        'RGPD compliant — cookies LSSI-CE, política de retención, AI Act Art. 52 compliance (transparencia IA)',
    ]
    for s in security:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('4.4 Stack Tecnológico', level=2)
    add_styled_table(doc,
        ['Capa', 'Tecnología'],
        [
            ['Backend', 'Python 3.12, FastAPI, framework multi-agente propietario'],
            ['Frontend', 'React 18, TypeScript, Vite 5, PWA'],
            ['IA (chat)', 'OpenAI GPT (últimas generaciones), embeddings de alta dimensión'],
            ['IA (OCR facturas)', 'Gemini 3 Flash Vision (~0,0003 $/factura, 33x más barato que Azure DI)'],
            ['Base de datos', 'Base de datos distribuida, edge-first'],
            ['Cache', 'Cache en memoria + cache semántico vectorial'],
            ['Pagos', 'Stripe (suscripciones + portal gestión)'],
            ['PDF Export', 'Generación de informes IRPF en PDF'],
            ['Email', 'Resend (envío informes a asesor)'],
            ['Despliegue', 'Railway (backend + frontend)'],
            ['Seguridad IA', 'Moderación IA multicapa (Llama Guard 4 + Prompt Guard 2)'],
        ]
    )

    doc.add_heading('4.5 Simulador IRPF Completo', level=2)
    doc.add_paragraph(
        'Impuestify incluye un motor de cálculo IRPF que cubre todos los tramos '
        'estatales y autonómicos de los 21 territorios soportados. El simulador se '
        'organiza en dos fases y extiende el régimen común con soporte foral completo:'
    )
    phase_items = [
        'Fase 1: Rendimientos del trabajo, del ahorro y de inmuebles. '
        'Reducciones por planes de pensiones. Deducciones: hipoteca pre-2013, '
        'maternidad, familia numerosa, donativos.',
        'Fase 2: Tributación conjunta, alquiler habitual pre-2015, rentas imputadas '
        'por inmuebles a disposición.',
        'Motor IRPF foral completo: escalas propias de Araba, Bizkaia, Gipuzkoa y Navarra '
        'con mínimos como deducción de cuota.',
        'Deducción 60% cuota íntegra para residentes en Ceuta y Melilla (Art. 68.4 LIRPF).',
        'Perfil fiscal adaptativo: formulario dinámico con campos específicos por CCAA y régimen fiscal.',
        'REST endpoint /api/irpf/estimate con debounce y AbortController para '
        'estimación en tiempo real sin saturar el servidor.',
        'Guía Fiscal: wizard interactivo de 7 pasos con LiveEstimatorBar que muestra '
        'la estimación IRPF actualizada en cada paso del proceso.',
    ]
    for item in phase_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.6 Clasificador de Facturas con IA (Phase 3 — NEW)', level=2)
    doc.add_paragraph(
        'Impuestify incorpora un clasificador de facturas basado en OCR con Gemini 3 Flash Vision '
        'que extrae automáticamente los campos de cada factura (emisor, CIF, fecha, base imponible, '
        'IVA, total) y la clasifica en la cuenta PGC correspondiente. El coste por factura es de '
        '~0,0003 $ (33x más barato que Azure Document Intelligence).'
    )
    invoice_items = [
        'OCR multimodal: procesa PDF, imagen y foto de factura con Gemini 3 Flash Vision',
        'Extracción estructurada: emisor, CIF, fecha, concepto, base, IVA, IRPF, total',
        'Clasificación PGC automática: asignación inteligente a cuenta contable (66 cuentas precargadas, grupos 1-7)',
        'Validación fiscal: verificación de CIF, tipos de IVA y coherencia de importes',
        '5 endpoints API dedicados: upload, list, detail, classify, bulk-process',
        'Coste ultra-bajo: ~0,0003 $/factura vs. ~0,01 $/factura en Azure DI (ahorro del 97 %)',
    ]
    for item in invoice_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.7 Contabilidad PGC (Phase 3 — NEW)', level=2)
    doc.add_paragraph(
        'Módulo de contabilidad de partida doble conforme al Plan General Contable (PGC) español. '
        'Permite a autónomos y pymes llevar su contabilidad sin necesidad de software adicional, '
        'con exportación directa para el Registro Mercantil.'
    )
    accounting_items = [
        'Libro Diario: registro cronológico de todos los asientos contables',
        'Libro Mayor: saldos por cuenta con detalle de movimientos',
        'Balance de Situación: activo, pasivo y patrimonio neto',
        'Cuenta de Pérdidas y Ganancias: ingresos, gastos y resultado del ejercicio',
        '66 cuentas PGC precargadas (grupos 1 a 7): inmovilizado, existencias, tesorería, proveedores, ventas, gastos, ingresos',
        'Asientos automáticos desde facturas clasificadas',
        'Exportación CSV/Excel compatible con Registro Mercantil',
        '5 endpoints API dedicados: journal, ledger, balance, pyg, export',
    ]
    for item in accounting_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # === 5. MERCADO OBJETIVO ===
    doc.add_heading('5. Mercado Objetivo', level=1)

    doc.add_heading('5.1 TAM / SAM / SOM', level=2)
    add_styled_table(doc,
        ['Segmento', 'Tamaño', 'Descripción'],
        [
            ['TAM', '21,8M declarantes IRPF', 'Total de declaraciones de renta anuales en España'],
            ['SAM', '5,0M declarantes digitales', 'Usuarios que usan herramientas digitales para su fiscalidad'],
            ['SOM Fase 1', '1,6M declarantes forales', 'Desatendidos por TaxDown y competencia — mercado abierto'],
            ['SOM Fase 2', '3,4M autónomos', 'Autónomos que pagan 50-150 €/mes a gestores humanos'],
        ],
        highlight_col=1
    )

    doc.add_heading('5.2 Segmentos prioritarios', level=2)

    segments = [
        ('Residentes forales (País Vasco + Navarra)',
         '1,6M declarantes que TaxDown rechaza explícitamente. Ningún competidor digital los atiende. '
         'Impuestify es la ÚNICA opción con IA. Captación estimada: 5 % = 80.000 usuarios potenciales.'),
        ('Creadores de contenido',
         '~285K creadores activos en España (YouTubers, streamers, influencers). Necesitan '
         'IVA por plataforma, Modelo 349, DAC7 y withholding tax. Ningún competidor cubre esto con IA. '
         'Plan Creator a 49 €/mes.'),
        ('Autónomos con complejidad media',
         '3,4M autónomos que gestionan IVA trimestral, IRPF fraccionado y cuota RETA. '
         'Hoy pagan 50-150 €/mes a gestores humanos con respuestas en 24-48 h. '
         'Impuestify ofrece respuestas en segundos a 39 €/mes.'),
        ('Particulares tecnológicos',
         'Trabajadores por cuenta ajena habituados a IA (ChatGPT, Gemini) que quieren un asistente '
         'fiscal inteligente, no un formulario. Plan Particular a 5 €/mes.'),
    ]
    for title, desc in segments:
        p = doc.add_paragraph()
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # === 6. ANÁLISIS COMPETITIVO ===
    doc.add_heading('6. Análisis Competitivo', level=1)

    doc.add_heading('6.1 Mapa de competidores', level=2)
    add_styled_table(doc,
        ['Plataforma', 'Tipo', 'Precio/mes', 'IA Real', 'Clasificador Facturas', 'Contabilidad PGC', 'Forales', 'Territorios'],
        [
            ['TaxDown', 'Asistente + asesor', '0-72 €', 'Básico (chatbot)', 'NO', 'NO', 'NO', 'Régimen común'],
            ['Declarando', 'Asesor autónomos', '29-99 €', 'No', 'SÍ (caro)', 'Básica', 'NO', 'Limitado'],
            ['Quipu', 'Facturación + contab.', '14,99-39,99 €', 'No', 'SÍ', 'SÍ (básica)', 'NO', 'Sin cobertura fiscal'],
            ['Holded', 'ERP completo', '29-99 €', 'No', 'SÍ', 'SÍ', 'NO', 'Sin simulador IRPF'],
            ['Sage', 'Contabilidad legacy', '30-100 €', 'No', 'No', 'SÍ', 'Parcial', 'Genérico'],
            ['Taxfix', 'Asesor humano', '~48 €', 'No', 'NO', 'NO', 'NO', 'Régimen común'],
            ['Gestor presencial', 'Humano', '80-150 €', 'No', 'Manual', 'SÍ', 'Depende', 'Variable'],
            ['IMPUESTIFY', 'Asistente IA full', '5-49 €', 'SÍ (multi-agente + Gemini Vision)', 'SÍ (IA, 0,0003 $/fact.)', 'SÍ (PGC completa)', 'SÍ (ÚNICO)', '21 territorios'],
        ],
        highlight_col=0
    )

    doc.add_heading('6.2 Ventajas competitivas de Impuestify', level=2)
    add_styled_table(doc,
        ['Ventaja', 'Detalle', 'Impacto'],
        [
            ['Cobertura foral completa', 'Único en cubrir PV + Navarra con IA', 'DIFERENCIADOR CRÍTICO'],
            ['Calculadora Neto 5 regímenes', 'ÚNICA calculadora en España que auto-detecta IVA/IGIC/IPSI, escala foral y deducción 60% Ceuta/Melilla por CCAA. Cuota SS por ingresos reales (15 tramos RDL 13/2022)', 'DIFERENCIADOR CRÍTICO'],
            ['Guía fiscal adaptativa', 'Wizard que adapta pasos por rol: particular (7), creator con plataformas (8), autónomo (8). Resultado con obligaciones específicas por perfil', 'DIFERENCIADOR'],
            ['Sistema multi-agente', '5 agentes especializados vs chatbot genérico', 'Calidad superior'],
            ['RAG sobre 439+ archivos oficiales', 'Respuestas con citas legales verificables', 'Confianza'],
            ['Análisis automático nóminas', 'Extracción PDF + proyección IRPF con extracción de 35 campos', 'Exclusivo'],
            ['Análisis notificaciones AEAT', 'IA en segundos vs humanos en 24h', 'Exclusivo'],
            ['~1.000 deducciones', '21 territorios cubiertos al 100 %, XSD Modelo 100 AEAT', 'Amplitud'],
            ['ISD + IPSI calculadoras', 'Sucesiones, donaciones y impuesto Ceuta/Melilla', 'Exclusivo'],
            ['Calendario fiscal personalizado', 'Plazos trimestrales por perfil con notificaciones push + email', 'Diferenciador'],
            ['Adjuntar docs en chat', 'Upload directo sin workspace, anonimización PII', 'UX superior'],
            ['Guardrails seguridad IA', 'Sistema de seguridad multicapa, anti-PII, anti-injection', 'Confianza regulatoria'],
            ['Calculadora retenciones IRPF', 'Algoritmo oficial AEAT, gratuita (lead magnet SEO)', 'Captación'],
            ['Compartir conversaciones', 'Enlaces públicos con anonimización PII automática', 'Viralidad'],
            ['Declaración conjunta vs. individual', 'Comparador de 4 escenarios con recomendación', 'Exclusivo'],
            ['GP inmuebles y plusvalía', 'Ganancias patrimoniales, renta imputada, ISD 21 CCAA', 'Amplitud'],
            ['Clasificador facturas IA', 'Gemini 3 Flash Vision: OCR + clasificación PGC automática a 0,0003 $/factura (33x más barato que Azure DI). Declarando cobra 29-99 EUR/mes', 'DIFERENCIADOR CRÍTICO'],
            ['Contabilidad PGC integrada', 'Libro Diario, Mayor, Balance, PyG + export Registro Mercantil. Quipu/Holded son standalone sin IA fiscal', 'DIFERENCIADOR'],
            ['Plataforma todo-en-uno', 'IRPF + IVA + facturas + contabilidad + deducciones + IA en un solo producto a 39 EUR/mes. Competidores requieren 2-3 herramientas', 'DIFERENCIADOR CRÍTICO'],
            ['Coste marginal casi cero', 'Sin asesores humanos, escalable', 'Margen superior'],
            ['Disponibilidad 24/7/365', 'No estacional como TaxDown', 'Engagement continuo'],
        ]
    )

    doc.add_heading('6.3 Gaps a cubrir (roadmap)', level=2)
    add_styled_table(doc,
        ['Gap', 'Competencia tiene', 'Plan Impuestify', 'Plazo'],
        [
            ['Presentación AEAT', 'TaxDown, Declarando', 'Colaborador Social AEAT', '6-12 meses'],
            ['Importación Cl@ve', 'TaxDown, Declarando', 'Requiere Colab. Social', '6-12 meses'],
            ['App móvil nativa', 'TaxDown, Declarando', 'React Native / PWA avanzada', '3-6 meses'],
            ['VeriFactu', 'A3, Sage', 'Obligatorio sociedades 01/01/2027', 'Q4 2026'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Gaps cerrados: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run(
        '(1) Motor de deducciones: TaxDown tiene ~250 (Rita), Impuestify ya tiene ~1.008 en 21 territorios. '
        '(2) Clasificador de facturas: Declarando cobra 29-99 EUR/mes, Impuestify lo incluye en el plan '
        'Autónomo (39 EUR/mes IVA incl.) con IA a 0,0003 $/factura. '
        '(3) Contabilidad PGC: Quipu/Holded son standalone sin IA fiscal; Impuestify integra contabilidad '
        'con simulador IRPF, deducciones y facturas en una sola plataforma.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'Nota: TaxDown tiene debilidades documentadas en Trustpilot (4.3/5, 6.539 reviews): '
        'discrepancias de cálculo de hasta 850 €, cobros duplicados, rotación de asesores, '
        'servicio lento, y prácticas de reviews cuestionadas por Trustpilot.'
    )
    run.italic = True
    run.font.size = Pt(9)

    doc.add_page_break()

    # === 7. MODELO DE NEGOCIO ===
    doc.add_heading('7. Modelo de Negocio', level=1)

    doc.add_heading('7.1 Planes de suscripción', level=2)
    add_styled_table(doc,
        ['Plan', 'Precio', 'Público', 'Incluye'],
        [
            ['Particular', '5 €/mes', 'Asalariados, pensionistas',
             'Chat IA ilimitado, IRPF 21 territorios, ~1.000 deducciones, análisis nóminas, '
             'notificaciones AEAT, workspace docs, export PDF, fuentes citadas, '
             'calendario fiscal, adjuntar docs en chat, ISD, compartir conversaciones, '
             'declaración conjunta vs. individual, múltiples pagadores'],
            ['Creator', '49 €/mes', 'YouTubers, streamers, influencers',
             'Todo lo de Particular + IVA por plataforma (YouTube, Twitch, TikTok...), '
             'Modelo 349 (intracomunitarias), DAC7, withholding tax W-8BEN, '
             'epígrafe IAE, CNAE 60.39, perfiles multirrol'],
            ['Autónomo', '39 €/mes (IVA incl.)', 'Autónomos y profesionales',
             'Todo lo de Particular + IVA Mod.303, IRPF Mod.130, cuota RETA, '
             'retenciones IRPF, deducciones autónomos, workspaces aislados, '
             'IPSI Ceuta/Melilla, Modelos 720/721, cobertura foral completa, '
             'clasificador de facturas con IA (OCR + PGC), contabilidad PGC completa '
             '(Libro Diario, Mayor, Balance, PyG), export CSV/Excel Registro Mercantil'],
        ],
        highlight_col=1
    )

    doc.add_heading('7.2 Proyección de ingresos (escenario conservador)', level=2)
    add_styled_table(doc,
        ['Métrica', 'Mes 6', 'Mes 12', 'Mes 24'],
        [
            ['Usuarios Particular', '500', '2.000', '8.000'],
            ['Usuarios Creator', '30', '50', '200'],
            ['Usuarios Autónomo', '100', '500', '2.000'],
            ['MRR Particular', '2.500 €', '10.000 €', '40.000 €'],
            ['MRR Creator', '1.470 €', '2.450 €', '9.800 €'],
            ['MRR Autónomo', '3.900 €', '19.500 €', '78.000 €'],
            ['MRR Total', '7.870 €', '31.950 €', '127.800 €'],
            ['ARR', '94.440 €', '383.400 €', '1.533.600 €'],
        ],
        highlight_col=3
    )

    doc.add_heading('7.3 Estructura de costes', level=2)
    add_styled_table(doc,
        ['Concepto', 'Coste mensual estimado', 'Notas'],
        [
            ['OpenAI API (GPT + embeddings)', '200-2.000 €', 'Escala con usuarios, caché semántica -30 %'],
            ['Gemini API (OCR facturas)', '5-50 €', '~0,0003 $/factura — 33x más barato que Azure DI'],
            ['Infraestructura (Railway + BD + caché)', '50-200 €', 'Serverless, escala automática'],
            ['Seguridad IA (moderación)', '0 €', 'Tier gratuito: 14.400 req/día'],
            ['Stripe (comisiones)', '1,4 % + 0,25 €/tx', 'Estándar Stripe UE'],
            ['Dominio + email', '~15 €', 'Fijo'],
        ]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        'Ventaja estructural: sin asesores humanos, el coste marginal por usuario tiende a cero. '
        'El modelo es inherentemente escalable frente a competidores como TaxDown (200+ asesores) '
        'o Declarando (equipos humanos por cliente).'
    )
    run.bold = True

    doc.add_page_break()

    # === 8. TRACCION Y METRICAS ===
    doc.add_heading('8. Tracción y Métricas', level=1)

    doc.add_heading('8.1 Estado actual del producto', level=2)
    metrics = [
        ('Producto', 'MVP completo, en producción en impuestify.com'),
        ('Usuarios registrados', '13 (fase alpha privada)'),
        ('Tests automatizados', '1.702 tests, 100 % passing'),
        ('Documentos RAG indexados', '456+ archivos fiscales oficiales'),
        ('Deducciones fiscales', '~1.008 activas en 21 territorios (21/21 cobertura)'),
        ('Agentes IA operativos', '6 agentes especializados + 1 Coordinator'),
        ('Herramientas IA', '15+ (IRPF, IVA, RETA, ISD, IPSI, deducciones, Mod.303/130, casillas, perfil fiscal, nóminas, notificaciones, retenciones, conjunta vs. individual, 720/721, clasificador facturas, contabilidad PGC)'),
        ('Casillas IRPF', '2.064 casillas Modelo 100 indexadas'),
        ('Simulador IRPF', 'Completo: 21 territorios, forales, Ceuta/Melilla, GP inmuebles, plusvalía, 2.º declarante'),
        ('Guía Fiscal', 'Wizard adaptativo por rol (7-8 pasos) con estimación IRPF en tiempo real'),
        ('Calculadora retenciones', 'Algoritmo oficial AEAT 2026, gratuita (lead magnet SEO)'),
        ('Compartir conversaciones', 'Enlaces públicos con anonimización PII automática'),
        ('Feedback', 'Widget flotante + valoración por respuesta + dashboard admin'),
        ('Documentos adjuntos', 'Upload en chat con clasificación automática y anonimización'),
        ('Clasificador Facturas', 'OCR con Gemini 3 Flash Vision + clasificación PGC automática (0,0003 $/factura)'),
        ('Contabilidad PGC', 'Libro Diario, Mayor, Balance, PyG. 66 cuentas PGC. Export CSV/Excel'),
        ('Seguridad', '13 capas de seguridad activas (rate limit, JWT, guardrails IA, PII, Turnstile, MFA)'),
        ('PWA', 'Instalable en móvil, offline-first para assets'),
        ('Pagos', 'Stripe integrado con 3 planes operativos (Particular, Creator, Autónomo)'),
    ]
    for title, value in metrics:
        p = doc.add_paragraph()
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(value)

    doc.add_heading('8.2 QA verificado en producción', level=2)
    doc.add_paragraph(
        '1.702 tests automatizados (backend) + E2E Playwright contra producción. '
        '25 sesiones de desarrollo completadas. Todas las funcionalidades principales operativas.'
    )

    doc.add_page_break()

    # === 9. ROADMAP ===
    doc.add_heading('9. Roadmap', level=1)

    add_styled_table(doc,
        ['Fase', 'Plazo', 'Objetivos', 'Estado'],
        [
            ['Fase 1: B2C Consolidación',
             'Q1-Q2 2026',
             'Motor IRPF 21 territorios, ~1.008 deducciones, 3 planes suscripción, '
             '13 capas seguridad, calc. retenciones, compartir conversaciones, '
             'feedback, GP inmuebles, ISD 21 CCAA, 720/721, 2.º declarante, auto-ingesta RAG',
             'COMPLETADO'],
            ['Fase 2: Verticales + Modelos',
             'Q2 2026',
             'Asesor modelos obligatorios por territorio, calculadora umbrales contables, '
             'vertical farmacias (backend + wizard + landing SEO), territory plugins',
             'COMPLETADO'],
            ['Fase 3: Facturas + Contabilidad',
             'Q2 2026',
             'Clasificador facturas IA (Gemini 3 Flash Vision, 0,0003 $/fact.), '
             'contabilidad PGC (Libro Diario, Mayor, Balance, PyG), '
             '66 cuentas PGC (grupos 1-7), 10 endpoints API, export CSV/Excel',
             'COMPLETADO'],
            ['Fase 4: B2B MVP',
             'Q3 2026',
             'Dashboard multi-cliente para asesorías, gestión de cartera, '
             'exportación masiva informes, multi-usuario por despacho, '
             'alertas proactivas de deducciones, API REST',
             'Planificado'],
            ['Fase 5: Compliance',
             'Q4 2026',
             'VeriFactu (obligatorio sociedades 01/01/2027), '
             'factura electrónica (Ley Crea y Crece), historial multi-ejercicio',
             'Planificado'],
            ['Fase 6: Crecimiento',
             '2027',
             'App móvil React Native, VeriFactu autónomos (01/07/2027), '
             'integración bancaria PSD2, ML predicción fiscal, B2B Enterprise white-label',
             'Planificado'],
        ]
    )

    doc.add_page_break()

    # === 10. EQUIPO ===
    doc.add_heading('10. Equipo', level=1)

    doc.add_paragraph(
        'Impuestify ha sido desarrollado íntegramente por un equipo lean con enfoque en IA '
        'y automatización. La arquitectura multi-agente y el RAG sobre documentación fiscal '
        'oficial representan más de 6 meses de desarrollo especializado.'
    )

    doc.add_heading('Perfiles a incorporar (con financiación)', level=2)
    roles = [
        ('CTO / Lead Developer', 'Arquitectura backend, escalabilidad, seguridad'),
        ('Growth / Marketing', 'Captación en nicho foral, SEO, partnerships'),
        ('Compliance / Legal', 'Trámites Colaborador Social AEAT, RGPD, licencias'),
        ('Customer Success', 'Onboarding autónomos, feedback, churn prevention'),
    ]
    for title, desc in roles:
        p = doc.add_paragraph()
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # === 11. FINANCIACION ===
    doc.add_heading('11. Financiación y Uso de Fondos', level=1)

    doc.add_heading('11.1 Ronda objetivo', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Pre-Seed: 150.000 - 300.000 €')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)

    doc.add_heading('11.2 Uso de fondos', level=2)
    add_styled_table(doc,
        ['Concepto', 'Porcentaje', 'Importe (sobre 200K €)', 'Detalle'],
        [
            ['Equipo técnico', '40 %', '80.000 €', 'CTO + developer senior (6 meses)'],
            ['Marketing y captación', '25 %', '50.000 €', 'SEO foral, paid acquisition, partnerships'],
            ['Compliance y legal', '15 %', '30.000 €', 'Colaborador Social AEAT, certificados, RGPD'],
            ['Infraestructura y ops', '10 %', '20.000 €', 'Servidores, APIs IA, herramientas'],
            ['Buffer operativo', '10 %', '20.000 €', 'Reserva 3 meses'],
        ],
        highlight_col=2
    )

    doc.add_heading('11.3 Hitos con la financiación', level=2)
    milestones = [
        'Mes 3: 500 usuarios activos, SEO posicionado en forales',
        'Mes 6: 1.000 usuarios, presentación AEAT en trámite',
        'Mes 9: App móvil, 150+ deducciones, primeros autónomos de pago',
        'Mes 12: Colaborador Social operativo, 2.500 usuarios, ARR ~383K €',
    ]
    for m in milestones:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()

    # Final CTA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('impuestify.com')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Contacto: fernando.prada@proton.me')
    run.font.size = Pt(11)

    # Save
    output_path = os.path.join(OUTPUT_DIR, "Impuestify_Business_Plan_2026.docx")
    doc.save(output_path)
    print(f"Business Plan saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
