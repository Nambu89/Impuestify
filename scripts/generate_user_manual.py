"""
Generate Impuestify User Manual as professional Word document.
Separate sections for Particulares, Creadores de Contenido and Autónomos.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "plans")


def set_cell_shading(cell, color_hex):
    shading = cell._tc.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1a56db')
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'f0f5ff')
    return table


def add_step(doc, number, title, description):
    p = doc.add_paragraph()
    run = p.add_run(f'Paso {number}: {title}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(26, 86, 219)
    doc.add_paragraph(description)


def build_manual():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Montserrat'
    style.font.size = Pt(11)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Montserrat'
        hs.font.color.rgb = RGBColor(26, 86, 219)

    for ls in ['List Bullet', 'List Number']:
        if ls in doc.styles:
            doc.styles[ls].font.name = 'Montserrat'
            doc.styles[ls].font.size = Pt(11)

    # =========================================================================
    # COVER
    # =========================================================================
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('IMPUESTIFY')
    run.font.size = Pt(38)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Manual de Usuario')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Guía completa para Particulares, Creadores de Contenido y Autónomos')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(26, 86, 219)

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run('Versión 3.0 — Abril 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    doc.add_heading('Índice', level=1)
    toc = [
        'PARTE I — Para todos los usuarios',
        '  1. Qué es Impuestify',
        '  2. Crear tu cuenta',
        '  3. Elegir tu plan',
        '  4. Tu primer chat con la IA',
        '  5. Perfil fiscal',
        '  6. Deducciones personalizadas',
        '  7. Informe IRPF en PDF',
        '  8. Guía Fiscal interactiva',
        '  8a. Calculadora de Sueldo Neto',
        '  8b. Calculadora de Retenciones IRPF',
        '  8c. Calendario fiscal',
        '  8d. Compartir conversaciones',
        '  8e. Declaración conjunta vs. individual',
        '  8f. Múltiples pagadores',
        '',
        'PARTE II — Plan Particular',
        '  9. Consultas IRPF por comunidad autónoma',
        '  10. Análisis de nóminas',
        '  11. Notificaciones de Hacienda',
        '',
        'PARTE III — Plan Creador de Contenido',
        '  12. Tu plan como creador',
        '',
        'PARTE IV — Plan Autónomo',
        '  13. Cuota de autónomos (RETA)',
        '  14. IVA trimestral (Modelo 303)',
        '  15. Pago fraccionado IRPF (Modelo 130)',
        '  15b. IPSI (Ceuta y Melilla)',
        '  15c. ISD (Sucesiones y Donaciones)',
        '  16. Retenciones IRPF en facturas',
        '  17. Workspaces de documentos',
        '  17b. Modelos 720 y 721 (bienes en el extranjero)',
        '  17c. Clasificador de Facturas',
        '  17d. Contabilidad y Libros',
        '',
        'PARTE V — Referencia',
        '  18. Territorios forales',
        '  19. Preguntas frecuentes',
        '  20. Soporte y contacto',
    ]
    for item in toc:
        if item == '':
            doc.add_paragraph()
            continue
        p = doc.add_paragraph(item)
        if item.startswith('PARTE'):
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(26, 86, 219)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # =========================================================================
    # PARTE I — PARA TODOS
    # =========================================================================
    p = doc.add_paragraph()
    run = p.add_run('PARTE I')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Para todos los usuarios')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # --- 1. Qué es Impuestify ---
    doc.add_heading('1. Qué es Impuestify', level=1)
    doc.add_paragraph(
        'Impuestify es tu asistente fiscal con Inteligencia Artificial. Puedes hacerle cualquier '
        'pregunta sobre impuestos en España y te responderá en segundos, citando fuentes legales '
        'oficiales. Es como tener un asesor fiscal disponible 24 horas al día, 7 días a la semana.'
    )

    doc.add_heading('Qué puede hacer por ti', level=2)
    capabilities = [
        'Calcular tu IRPF en cualquier comunidad autónoma (incluyendo País Vasco y Navarra)',
        'Descubrir deducciones fiscales que te corresponden entre más de 1.000 opciones en 21 territorios',
        'Adjuntar documentos directamente en el chat para análisis instantáneo',
        'Calcular el Impuesto de Sucesiones y Donaciones (ISD) en todas las comunidades',
        'Consultar tu calendario fiscal personalizado con fechas límite',
        'Actualizar tu perfil fiscal automáticamente desde la conversación',
        'Analizar tus nóminas en PDF y detectar errores de retención',
        'Interpretar notificaciones de Hacienda y decirte qué hacer',
        'Generar un informe IRPF exportable en PDF',
        'Comparar declaración conjunta vs. individual y elegir la más ventajosa',
        'Calcular retenciones IRPF con el algoritmo oficial de la AEAT',
        'Compartir conversaciones con tu asesor mediante enlaces públicos seguros',
        'Para autónomos: calcular IVA, cuota RETA, pagos fraccionados y retenciones',
        'Para creadores de contenido: IVA por plataforma, Modelo 349, DAC7, epígrafe IAE',
        'Clasificar facturas automáticamente con OCR y asignar cuentas del Plan General Contable',
        'Generar libros contables (Diario, Mayor, Balance de Sumas y Saldos, Pérdidas y Ganancias) y exportar a CSV/Excel',
    ]
    for c in capabilities:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_heading('Qué NO puede hacer (todavía)', level=2)
    limitations = [
        'Presentar tu declaración de la renta directamente a Hacienda (en desarrollo)',
        'Importar datos automáticamente con Cl@ve PIN (en desarrollo)',
        'Sustituir a un asesor en situaciones muy complejas (inversiones internacionales, herencias)',
    ]
    for l in limitations:
        doc.add_paragraph(l, style='List Bullet')

    # --- 2. Crear cuenta ---
    doc.add_heading('2. Crear tu cuenta', level=1)
    add_step(doc, 1, 'Accede a impuestify.com', 'Abre tu navegador y ve a impuestify.com.')
    add_step(doc, 2, 'Pulsa "Empezar Ahora"', 'En la página principal, haz clic en el botón azul.')
    add_step(doc, 3, 'Rellena tus datos',
             'Introduce tu nombre, email y una contraseña segura. '
             'Tu email será tu identificador único.')
    add_step(doc, 4, 'Selecciona tu comunidad autónoma',
             'Durante el registro deberás indicar tu comunidad autónoma. Este dato es obligatorio '
             'y determina los tramos IRPF, las deducciones territoriales y el calendario fiscal '
             'que verás en la aplicación.')
    add_step(doc, 5, 'Completa el control de seguridad',
             'Verás un control de seguridad Cloudflare Turnstile (un clic o resolución automática). '
             'Este paso protege la plataforma contra registros automáticos.')
    add_step(doc, 6, 'Confirma tu cuenta',
             'Recibirás un email de confirmación. Haz clic en el enlace para activar tu cuenta.')

    # --- 3. Elegir plan ---
    doc.add_heading('3. Elegir tu plan', level=1)
    doc.add_paragraph('Impuestify ofrece tres planes adaptados a tu situación:')

    add_table(doc,
        ['Característica', 'Particular (5 €/mes)', 'Creador (49 €/mes)', 'Autónomo (39 €/mes IVA incl.)'],
        [
            ['Chat IA fiscal ilimitado', 'SÍ', 'SÍ', 'SÍ'],
            ['Cálculo IRPF por CCAA', 'SÍ', 'SÍ', 'SÍ'],
            ['1.000+ deducciones personalizadas', 'SÍ', 'SÍ', 'SÍ'],
            ['Análisis de nóminas (PDF)', 'SÍ', 'SÍ', 'SÍ'],
            ['Adjuntar documentos en el chat', 'SÍ', 'SÍ', 'SÍ'],
            ['Calendario fiscal personalizado', 'SÍ', 'SÍ', 'SÍ'],
            ['Cálculo ISD (Sucesiones)', 'SÍ', 'SÍ', 'SÍ'],
            ['Notificaciones AEAT', 'SÍ', 'SÍ', 'SÍ'],
            ['Workspace de documentos', 'SÍ', 'SÍ', 'SÍ'],
            ['Informe IRPF en PDF', 'SÍ', 'SÍ', 'SÍ'],
            ['Perfil fiscal adaptativo por CCAA', 'SÍ', 'SÍ', 'SÍ'],
            ['Compartir conversaciones', 'SÍ', 'SÍ', 'SÍ'],
            ['Declaración conjunta vs. individual', 'SÍ', 'SÍ', 'SÍ'],
            ['IVA por plataforma (YouTube, Twitch...)', 'NO', 'SÍ', 'NO'],
            ['Modelo 349 (intracomunitarias)', 'NO', 'SÍ', 'SÍ'],
            ['DAC7 y withholding tax', 'NO', 'SÍ', 'NO'],
            ['Perfiles multirrol', 'NO', 'SÍ', 'NO'],
            ['Cuota autónomos (RETA)', 'NO', 'NO', 'SÍ'],
            ['IVA trimestral (Mod. 303)', 'NO', 'NO', 'SÍ'],
            ['Pago fraccionado (Mod. 130)', 'NO', 'NO', 'SÍ'],
            ['IPSI Ceuta/Melilla', 'NO', 'NO', 'SÍ'],
            ['Retenciones IRPF en facturas', 'NO', 'NO', 'SÍ'],
            ['Modelos 720/721 (bienes extranjero)', 'NO', 'NO', 'SÍ'],
            ['Clasificador de facturas (OCR + PGC)', 'NO', 'SÍ', 'SÍ'],
            ['Contabilidad y libros (Diario, Mayor...)', 'NO', 'SÍ', 'SÍ'],
            ['Cobertura foral completa', 'SÍ', 'SÍ', 'SÍ'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Consejo: ')
    run.bold = True
    p.add_run(
        'Si eres trabajador por cuenta ajena, el Plan Particular es perfecto para ti. '
        'Si eres youtuber, streamer, instagramer o cualquier tipo de creador de contenido, '
        'el Plan Creador incluye IVA por plataforma, DAC7 y Modelo 349. '
        'Si eres autónomo o profesional por cuenta propia, el Plan Autónomo '
        'te da acceso a las calculadoras de IVA, RETA y pagos fraccionados.'
    )

    doc.add_page_break()

    # --- 4. Tu primer chat ---
    doc.add_heading('4. Tu primer chat con la IA', level=1)
    doc.add_paragraph(
        'Una vez dentro, verás el chat de Impuestify. Puedes escribir cualquier pregunta '
        'fiscal en lenguaje natural, como si hablaras con un asesor.'
    )

    doc.add_heading('Ejemplos de preguntas', level=2)
    examples = [
        '"Cuánto IRPF pago si gano 40.000 euros en Madrid?"',
        '"Que deducciones me corresponden si tengo hijos y vivo en Cataluña?"',
        '"He recibido una notificación de Hacienda, qué hago?"',
        '"Cuánto es mi cuota de autónomos si gano 2.000 euros al mes?"',
        '"Cómo calculo el IVA del trimestre?"',
    ]
    for e in examples:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_heading('Cómo funciona la respuesta', level=2)
    doc.add_paragraph(
        'Cuando envías una pregunta, verás un timeline de pasos que muestra cómo piensa la IA:'
    )
    steps = [
        ('Pensando...', 'La IA analiza tu pregunta y decide qué herramientas necesita.'),
        ('Buscando normativa...', 'Consulta entre 439+ documentos oficiales.'),
        ('Calculando...', 'Ejecuta las calculadoras necesarias (IRPF, deducciones, etc.).'),
        ('Respuesta', 'Te muestra el resultado con explicación y fuentes citadas.'),
    ]
    for title, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(f'{title} ')
        run.bold = True
        run.font.color.rgb = RGBColor(26, 86, 219)
        p.add_run(desc)

    doc.add_heading('Adjuntar documentos', level=2)
    doc.add_paragraph(
        'Junto al campo de texto verás un icono de clip. Pulsa para adjuntar un documento '
        '(PDF, JPG, PNG). El sistema lo analiza automáticamente, extrae los datos relevantes '
        'y los usa como contexto para la conversación. Los documentos adjuntos son temporales: '
        'se eliminan al cerrar el navegador.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Ejemplo: ')
    run.bold = True
    p.add_run('Adjunta tu nómina y pregunta: "¿Es correcta mi retención IRPF?"')

    # --- 5. Perfil fiscal ---
    doc.add_heading('5. Perfil fiscal', level=1)
    doc.add_paragraph(
        'Tu perfil fiscal permite a la IA personalizar las respuestas. '
        'Ve a Ajustes > Perfil Fiscal para configurarlo. '
        'El perfil se puede rellenar manualmente en Ajustes o automáticamente desde la Guía Fiscal.'
    )

    doc.add_paragraph(
        'Tu perfil fiscal se adapta dinámicamente según tu comunidad autónoma. '
        'Si vives en un territorio foral (Araba, Bizkaia, Gipuzkoa, Navarra) verás campos '
        'específicos como EPSV, cuenta vivienda foral, etc.'
    )

    doc.add_paragraph(
        'La IA puede actualizar tu perfil automáticamente. Sube una nómina y pide: '
        '"Guarda estos datos en mi perfil fiscal". La IA extraerá los datos y los guardará.'
    )

    doc.add_heading('Datos del perfil', level=2)
    profile_fields = [
        ('Comunidad autónoma', 'Donde resides. Determina los tramos IRPF y deducciones territoriales.'),
        ('Situación laboral', 'Particular (cuenta ajena) o Autónomo (cuenta propia).'),
        ('Deducciones y situación familiar',
         'Planes de pensiones, hipoteca pre-2013, maternidad, familia numerosa, donativos a ONGs, '
         'tributación conjunta, alquiler pre-2015, rentas imputadas por segundas viviendas.'),
        ('Discapacidad', 'Grado de discapacidad del contribuyente y de familiares a cargo.'),
        ('Vivienda', 'Tipo de situación: alquiler, propiedad, rehabilitación, zona rural.'),
        ('Sostenibilidad', 'Vehículo eléctrico, instalación de paneles solares, obras de eficiencia energética.'),
        ('Donaciones', 'Aportaciones a entidades autonómicas, investigación, patrimonio histórico.'),
        ('Para autónomos', 'Epígrafe IAE, base de cotización, tipo de IVA, retención IRPF, '
         'régimen SS, gastos deducibles, amortizaciones...'),
    ]
    for title, desc in profile_fields:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)

    p = doc.add_paragraph()
    run = p.add_run('Importante: ')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 50, 50)
    p.add_run('La IA recuerda tu perfil entre sesiones. Si cambias de comunidad '
              'autónoma o de situación laboral, actualízalo en Ajustes.')

    # --- 6. Deducciones ---
    doc.add_heading('6. Deducciones personalizadas', level=1)
    doc.add_paragraph(
        'Impuestify tiene un motor de más de 1.000 deducciones fiscales en 21 territorios: '
        'estatales, autonómicas y forales. Es la cobertura más completa del mercado. '
        'Cuando preguntas por deducciones, la IA:'
    )

    deduction_steps = [
        'Identifica tu comunidad autónoma (del perfil o preguntando)',
        'Filtra las deducciones aplicables a tu territorio',
        'Te hace preguntas de elegibilidad (hijos, alquiler, discapacidad, etc.)',
        'Calcula el ahorro estimado para cada deducción',
        'Muestra las deducciones en tarjetas visuales con el detalle legal',
    ]
    for s in deduction_steps:
        doc.add_paragraph(s, style='List Number')

    p = doc.add_paragraph()
    run = p.add_run('Ejemplo: ')
    run.bold = True
    p.add_run(
        'Un residente en Araba (País Vasco) con hijos puede acceder a deducciones forales '
        'como la deducción por descendientes de 734,80 EUR por hijo, la deducción por alquiler '
        'del 20% (max 1.600 EUR), o la deducción por compra de vivienda del 18%. '
        'Estas deducciones NO existen en el régimen común y ningún otro asistente digital las cubre.'
    )

    doc.add_paragraph(
        'Para residentes forales (Araba, Bizkaia, Gipuzkoa, Navarra) el sistema aplica '
        'automáticamente solo las deducciones forales, ya que tienen un IRPF completamente '
        'independiente del régimen estatal.'
    )

    # --- 7. Informe PDF ---
    doc.add_heading('7. Informe IRPF en PDF', level=1)
    doc.add_paragraph(
        'Después de una simulación IRPF, puedes exportar un informe completo en PDF con:'
    )
    pdf_contents = [
        'Resumen de ingresos y gastos deducibles',
        'Cálculo detallado por tramos IRPF (estatal + autonómico)',
        'Deducciones aplicables con importes',
        'Cuota resultante y tipo efectivo',
        'Referencias legales de cada cálculo',
    ]
    for c in pdf_contents:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph(
        'Puedes descargar el PDF o enviarlo directamente por email a tu asesor '
        'desde el botón "Enviar a asesor" que aparece junto al informe.'
    )

    # --- 8. Guía Fiscal interactiva ---
    doc.add_heading('8. Guía Fiscal interactiva (adaptativa por rol)', level=1)
    doc.add_paragraph(
        'La Guía Fiscal es un asistente paso a paso que adapta su contenido según tu plan de suscripción. '
        'Accede desde el menú principal > "Guía Fiscal".'
    )

    doc.add_heading('Plan Particular (7 pasos)', level=2)
    doc.add_paragraph(
        'Si tu plan es Particular, la guía muestra 7 pasos simplificados sin la sección de actividad económica:'
    )
    add_step(doc, 1, 'Datos personales',
             'Comunidad autónoma, edad, tributación conjunta. Si vives en Ceuta/Melilla, bonificación 60% automática.')
    add_step(doc, 2, 'Trabajo', 'Ingresos brutos, cotizaciones SS, retenciones IRPF.')
    add_step(doc, 3, 'Ahorro e inversiones', 'Intereses, dividendos, fondos, criptomonedas.')
    add_step(doc, 4, 'Inmuebles', 'Alquileres, deducción pre-2015, rentas imputadas.')
    add_step(doc, 5, 'Familia', 'Descendientes, ascendientes, discapacidad, maternidad.')
    add_step(doc, 6, 'Deducciones', 'Planes de pensiones, hipoteca pre-2013, donativos, deducciones autonómicas.')
    add_step(doc, 7, 'Resultado', 'Cuota líquida, retenciones, resultado (a pagar o devolver), tipo medio efectivo.')

    doc.add_heading('Plan Creator (8 pasos)', level=2)
    doc.add_paragraph(
        'Si tu plan es Creator, la guía incluye un paso dedicado "Actividad como creador" con:'
    )
    doc.add_paragraph(
        '• Grid de 10 plataformas (YouTube, Twitch, TikTok, Instagram, OnlyFans, Patreon, Substack, sponsors, merch, otros)\n'
        '• Selector de epígrafe IAE (8690, 9020, 6010.1, 961.1)\n'
        '• Gastos deducibles de creador (equipo, software, coworking, transporte, formación)\n'
        '• Información sobre IVA intracomunitario (Google Ireland, Amazon Luxembourg, Meta Ireland)\n'
        '• Withholding tax W-8BEN (retención plataformas USA)\n'
        '• Toggle Modelo 349 (operaciones intracomunitarias)\n'
        '• Resultado con obligaciones: Modelo 349, Modelo 130, DAC7, epígrafe IAE'
    )

    doc.add_heading('Plan Autónomo (8 pasos)', level=2)
    doc.add_paragraph(
        'Si tu plan es Autónomo, la guía incluye un paso dedicado "Actividad económica" con ingresos, '
        'gastos, cuota autónomos, amortizaciones, régimen de estimación (directa simplificada/normal/objetiva), '
        'retenciones y pagos fraccionados. El resultado muestra obligaciones: Modelo 130, Modelo 303, cuota RETA.'
    )

    doc.add_paragraph(
        'Mientras completas los pasos, una barra de estimación en tiempo real te muestra el resultado '
        'parcial. Verde si te sale a devolver, rojo si te sale a pagar.'
    )

    doc.add_heading('8a. Calculadora Sueldo Neto (NUEVO)', level=1)
    doc.add_paragraph(
        'La Calculadora de Sueldo Neto responde a la pregunta más frecuente de los autónomos: '
        '"¿Cuánto me queda limpio de lo que facturo?". Accede desde el menú > "Calculadora Neto" '
        'o directamente en /calculadora-neto.'
    )
    doc.add_paragraph(
        'Introduce tu facturación mensual bruta y obtén un desglose inmediato: IVA/IGIC/IPSI repercutido, '
        'retención IRPF, cuota de autónomos, gastos deducibles y tu neto real mensual y anual.'
    )
    doc.add_heading('5 regímenes fiscales automáticos', level=2)
    doc.add_paragraph(
        'La calculadora detecta automáticamente tu régimen fiscal según tu comunidad autónoma:\n\n'
        '• Madrid, Andalucía y resto de régimen común: IVA 21%, escala IRPF estatal + autonómica\n'
        '• Canarias: IGIC 7% (en vez de IVA 21%), misma escala IRPF\n'
        '• Ceuta y Melilla: IPSI 4% + deducción del 60% sobre la cuota IRPF (Art. 68.4 LIRPF)\n'
        '• País Vasco (Araba, Bizkaia, Gipuzkoa): IVA 21% pero escala IRPF foral propia (7 tramos)\n'
        '• Navarra: escala IRPF foral propia (11 tramos)'
    )
    doc.add_paragraph(
        'La cuota de Seguridad Social se calcula automáticamente según tus ingresos reales '
        '(sistema de cotización por ingresos reales, RDL 13/2022, 15 tramos). '
        'También puedes introducir tu cuota manualmente si la conoces.'
    )
    doc.add_paragraph(
        'Nota: esta estimación es orientativa. El resultado real depende de tu situación personal, '
        'familiar y de tu comunidad autónoma. Consulta con un asesor fiscal para una declaración precisa.'
    )

    # --- 8b. Calculadora de Retenciones IRPF ---
    doc.add_heading('8b. Calculadora de Retenciones IRPF (NUEVO)', level=1)
    doc.add_paragraph(
        'La Calculadora de Retenciones IRPF aplica el algoritmo oficial de la AEAT para 2026. '
        'Accede desde el menú > "Calculadora Retenciones" o directamente en /calculadora-retenciones. '
        'Es una herramienta pública y gratuita (no requiere suscripción).'
    )
    doc.add_paragraph(
        'Introduce tu salario bruto anual, situación familiar y comunidad autónoma, '
        'y obtendrás el tipo de retención exacto que debe aplicar tu empresa en la nómina. '
        'Ideal para verificar si la retención que te aplican es correcta.'
    )
    doc.add_heading('Datos que se solicitan', level=2)
    retention_fields = [
        'Salario bruto anual',
        'Situación familiar (soltero/a, casado/a, número de hijos)',
        'Comunidad autónoma de residencia',
        'Discapacidad (si aplica)',
        'Contrato temporal o indefinido',
    ]
    for f in retention_fields:
        doc.add_paragraph(f, style='List Bullet')

    # --- 8c. Calendario fiscal ---
    doc.add_heading('8c. Calendario fiscal', level=1)
    doc.add_paragraph(
        'Impuestify incluye un calendario fiscal personalizado que te muestra tus próximas '
        'obligaciones tributarias. Accede desde el menú principal > "Calendario".'
    )

    doc.add_paragraph('El calendario se adapta a tu perfil:')
    calendar_items = [
        'Particular: plazos de declaración IRPF anual',
        'Autónomo: plazos trimestrales (Modelo 303 IVA, Modelo 130 IRPF, cuota RETA)',
        'Canarias: IGIC en lugar de IVA',
        'Ceuta/Melilla: IPSI en lugar de IVA',
    ]
    for item in calendar_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'Cada fecha límite incluye una descripción del trámite y los días restantes para presentarlo.'
    )

    # --- 8d. Compartir conversaciones ---
    doc.add_heading('8d. Compartir conversaciones (NUEVO)', level=1)
    doc.add_paragraph(
        'Puedes compartir cualquier conversación con tu asesor fiscal, familiar o gestor '
        'mediante un enlace público seguro. Pulsa el botón de compartir en la esquina '
        'superior del chat y se generará un enlace único.'
    )
    doc.add_paragraph(
        'Antes de compartir, Impuestify anonimiza automáticamente los datos sensibles: '
        'DNI/NIE, números IBAN, importes concretos y otros datos personales. '
        'La persona que reciba el enlace verá la conversación completa pero sin datos que '
        'permitan identificarte.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Nota de seguridad: ')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 50, 50)
    p.add_run('Los enlaces compartidos son de solo lectura. No permiten modificar '
              'la conversación ni acceder a tu cuenta.')

    # --- 8e. Declaración conjunta vs. individual ---
    doc.add_heading('8e. Declaración conjunta vs. individual (NUEVO)', level=1)
    doc.add_paragraph(
        'Si estás casado/a o formas una unidad familiar, Impuestify puede comparar '
        'automáticamente las cuatro posibles combinaciones de declaración:'
    )
    joint_scenarios = [
        'Ambos declarantes en individual',
        'Declaración conjunta',
        'Primer declarante individual + segundo conjunta',
        'Primer declarante conjunta + segundo individual',
    ]
    for s in joint_scenarios:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph(
        'El sistema calcula la cuota resultante para cada escenario y te recomienda '
        'la opción más ventajosa económicamente. Pregunta en el chat: '
        '"¿Me conviene más la declaración conjunta o la individual?"'
    )

    # --- 8f. Múltiples pagadores ---
    doc.add_heading('8f. Múltiples pagadores (NUEVO)', level=1)
    doc.add_paragraph(
        'Si durante el año has tenido más de un pagador (por ejemplo, has cambiado de empresa), '
        'Impuestify calcula automáticamente si estás obligado a presentar la declaración de la renta '
        'según el artículo 96 de la Ley del IRPF.'
    )
    doc.add_paragraph(
        'El límite general es de 22.000 € con un solo pagador, pero baja a 15.876 € si '
        'el segundo y siguientes pagadores suman más de 1.500 €. Indica tus pagadores en el chat '
        'y la IA te dirá si estás obligado y cuál sería tu resultado estimado.'
    )

    doc.add_page_break()

    # =========================================================================
    # PARTE II — PLAN PARTICULAR
    # =========================================================================
    p = doc.add_paragraph()
    run = p.add_run('PARTE II')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Plan Particular')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # --- 9. IRPF por CCAA ---
    doc.add_heading('9. Consultas IRPF por comunidad autónoma', level=1)
    doc.add_paragraph(
        'Pregunta por tu IRPF indicando tus ingresos y comunidad autónoma. '
        'La IA aplicará los tramos correctos (estatal + autonómico) y te mostrará:'
    )
    irpf_results = [
        'Base imponible general y del ahorro',
        'Cálculo tramo a tramo con tipos aplicables',
        'Cuota íntegra estatal y autonómica',
        'Deducciones aplicables',
        'Cuota líquida y tipo efectivo resultante',
    ]
    for r in irpf_results:
        doc.add_paragraph(r, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Ejemplo de pregunta: ')
    run.bold = True
    p.add_run('"Si gano 35.000 euros brutos al año en Valencia, cuánto IRPF pago?"')

    # --- 10. Nóminas ---
    doc.add_heading('10. Análisis de nóminas', level=1)
    doc.add_paragraph(
        'Adjunta tu nómina directamente en el chat con el botón de clip, o súbela a un workspace. '
        'La IA extraerá automáticamente:'
    )
    payslip_fields = [
        'Salario bruto y neto',
        'Retención IRPF (porcentaje y importe)',
        'Cotizaciones a la Seguridad Social',
        'Base de cotización SS',
        'Base IRPF',
        'Aportaciones de la empresa',
        'Grupo de cotización y categoría profesional',
        'Pagas extraordinarias',
        'Detalle de devengos (salario base, complementos, pluses)',
        'Proyección IRPF anual basada en la nómina',
    ]
    for f in payslip_fields:
        doc.add_paragraph(f, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Consejo: ')
    run.bold = True
    p.add_run('Si detectamos que tu retención IRPF es incorrecta respecto a tus ingresos anuales, '
              'te avisaremos automáticamente.')

    doc.add_paragraph(
        'La IA puede guardar los datos extraídos en tu perfil fiscal para futuras consultas. '
        'Pide: "Guarda estos datos en mi perfil fiscal."'
    )

    # --- 11. Notificaciones ---
    doc.add_heading('11. Notificaciones de Hacienda', level=1)
    doc.add_paragraph(
        'Si recibes una notificación de la AEAT (Agencia Tributaria), adjunta el PDF directamente '
        'en el chat o súbelo a un workspace. La IA extraerá:'
    )
    notif_fields = [
        'Tipo de notificación (requerimiento, liquidación, sanción, etc.)',
        'Plazo de respuesta',
        'Importe (si aplica)',
        'Referencia y número de expediente',
        'Acción requerida y pasos a seguir',
    ]
    for f in notif_fields:
        doc.add_paragraph(f, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Importante: ')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 50, 50)
    p.add_run('Las notificaciones de Hacienda tienen plazos legales. '
              'Si la notificación requiere acción urgente, te lo indicaremos claramente.')

    doc.add_page_break()

    # =========================================================================
    # PARTE III — PLAN CREADOR DE CONTENIDO
    # =========================================================================
    p = doc.add_paragraph()
    run = p.add_run('PARTE III')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Plan Creador de Contenido')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # --- 12. Tu plan como creador ---
    doc.add_heading('12. Tu plan como creador de contenido', level=1)
    doc.add_paragraph(
        'Si eres youtuber, streamer, instagramer, tiktoker, podcaster, blogger o cualquier '
        'tipo de creador de contenido digital, el Plan Creador (49 €/mes) está diseñado '
        'específicamente para ti. Cubre todas las obligaciones fiscales que afectan a los '
        'creadores en España.'
    )

    doc.add_heading('Plataformas cubiertas', level=2)
    doc.add_paragraph(
        'Impuestify conoce el tratamiento fiscal de las 10 plataformas principales y calcula '
        'el IVA correspondiente según la sede de cada una:'
    )
    platforms = [
        'YouTube (Google Ireland) — IVA intracomunitario',
        'Twitch (Amazon Luxembourg) — IVA intracomunitario',
        'TikTok — Ingresos del fondo de creadores',
        'Instagram/Facebook (Meta Ireland) — IVA intracomunitario',
        'OnlyFans — Retención y IVA según país',
        'Patreon — Suscripciones con IVA',
        'Substack — Newsletters de pago',
        'Patrocinadores directos — Facturación nacional',
        'Merchandising — Venta de productos propios',
        'Otros ingresos digitales',
    ]
    for pl in platforms:
        doc.add_paragraph(pl, style='List Bullet')

    doc.add_heading('Obligaciones específicas del creador', level=2)
    creator_obligations = [
        ('Epígrafe IAE', 'Te indicamos el epígrafe correcto según tu actividad: '
         '8690 (otros servicios), 9020 (espectáculos), 6010.1 (comercio) o 961.1 (producción).'),
        ('Modelo 349', 'Si facturas a plataformas con sede en la UE (Google Ireland, Meta Ireland, '
         'Amazon Luxembourg), debes presentar la declaración de operaciones intracomunitarias.'),
        ('DAC7', 'La Directiva DAC7 obliga a las plataformas digitales a informar a Hacienda '
         'sobre los ingresos de los creadores. Te explicamos qué datos comparten y cómo afecta a tu declaración.'),
        ('Withholding tax (W-8BEN)', 'Si recibes pagos de plataformas en EE. UU., la retención '
         'por defecto es del 30 %. Con el formulario W-8BEN se reduce al 15 % (convenio España-EE. UU.).'),
        ('CNAE 60.39', 'Código de actividad económica para creadores de contenido digital.'),
    ]
    for title, desc in creator_obligations:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('Perfiles multirrol', level=2)
    doc.add_paragraph(
        'Muchos creadores tienen también un empleo por cuenta ajena o son autónomos. '
        'Impuestify permite configurar roles adicionales (no exclusivos) para que tu perfil fiscal '
        'refleje todas tus fuentes de ingresos y el sistema aplique las reglas de cada rol.'
    )

    doc.add_page_break()

    # =========================================================================
    # PARTE IV — PLAN AUTÓNOMO
    # =========================================================================
    p = doc.add_paragraph()
    run = p.add_run('PARTE IV')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Plan Autónomo')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # --- 13. RETA ---
    doc.add_heading('13. Cuota de autónomos (RETA)', level=1)
    doc.add_paragraph(
        'Desde 2023, la cuota de autónomos se calcula en función de los rendimientos netos reales '
        '(sistema de tramos). Impuestify calcula tu cuota exacta preguntándote:'
    )
    reta_questions = [
        'Ingresos netos mensuales estimados',
        'Si estás en tarifa plana (primeros 12 meses: 80 EUR/mes)',
        'Tramo de cotización correspondiente',
    ]
    for q in reta_questions:
        doc.add_paragraph(q, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Ejemplo: ')
    run.bold = True
    p.add_run('"Soy autónomo y facturo 3.000 euros netos al mes, cuánto pago de cuota?"')

    doc.add_heading('Tabla de tramos RETA 2025', level=2)
    add_table(doc,
        ['Rendimientos netos/mes', 'Base mínima', 'Base máxima', 'Cuota mínima aprox.'],
        [
            ['Hasta 670 EUR', '735,29 EUR', '816,98 EUR', '~225 EUR'],
            ['670 - 900 EUR', '816,99 EUR', '900 EUR', '~250 EUR'],
            ['900 - 1.166,70 EUR', '872,55 EUR', '1.166,70 EUR', '~267 EUR'],
            ['1.166,70 - 1.300 EUR', '950,98 EUR', '1.300 EUR', '~291 EUR'],
            ['1.300 - 1.500 EUR', '960,78 EUR', '1.500 EUR', '~294 EUR'],
            ['1.500 - 1.700 EUR', '960,78 EUR', '1.700 EUR', '~294 EUR'],
            ['Más de 1.700 EUR', '1.143,79 EUR', '4.720,50 EUR', '~350+ EUR'],
        ]
    )

    doc.add_page_break()

    # --- 14. IVA ---
    doc.add_heading('14. IVA trimestral (Modelo 303)', level=1)
    doc.add_paragraph(
        'El Modelo 303 es la declaración trimestral de IVA. Impuestify te ayuda a calcularlo:'
    )

    add_step(doc, 1, 'Indica tus datos del trimestre',
             '"He facturado 15.000 EUR con IVA del 21% y tengo gastos deducibles por 3.000 EUR con IVA."')
    add_step(doc, 2, 'La IA calcula',
             'IVA repercutido (cobrado a clientes) - IVA soportado (pagado en gastos) = IVA a ingresar.')
    add_step(doc, 3, 'Resultado',
             'Te muestra el importe a pagar (o a compensar si el soportado es mayor) y el plazo de presentación.')

    doc.add_heading('Plazos de presentación', level=2)
    add_table(doc,
        ['Trimestre', 'Periodo', 'Plazo de presentación'],
        [
            ['1T', 'Enero - Marzo', '1 al 20 de abril'],
            ['2T', 'Abril - Junio', '1 al 20 de julio'],
            ['3T', 'Julio - Septiembre', '1 al 20 de octubre'],
            ['4T', 'Octubre - Diciembre', '1 al 30 de enero (año siguiente)'],
        ]
    )

    # --- 15. Modelo 130 ---
    doc.add_heading('15. Pago fraccionado IRPF (Modelo 130)', level=1)
    doc.add_paragraph(
        'Si estás en estimación directa, debes presentar pagos fraccionados trimestrales '
        'del IRPF (Modelo 130). Es un anticipo del 20% sobre el beneficio acumulado.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Fórmula: ')
    run.bold = True
    p.add_run('20% x (Ingresos acumulados - Gastos acumulados) - Pagos fraccionados anteriores')

    p = doc.add_paragraph()
    run = p.add_run('Ejemplo de pregunta: ')
    run.bold = True
    p.add_run('"He facturado 30.000 EUR en el primer semestre con gastos de 8.000 EUR. '
              'Cuánto debo pagar en el Modelo 130 del 2T?"')

    # --- 15b. IPSI Ceuta y Melilla ---
    doc.add_heading('15b. IPSI (Ceuta y Melilla)', level=1)
    doc.add_paragraph(
        'Si resides en Ceuta o Melilla, el IVA no aplica. En su lugar pagas el IPSI '
        '(Impuesto sobre la Producción, los Servicios y la Importación). '
        'Impuestify calcula el IPSI trimestral con los 6 tipos impositivos aplicables: '
        '0,5%, 1%, 2%, 4%, 8% y 10%.'
    )

    # --- 15c. ISD ---
    doc.add_heading('15c. ISD (Sucesiones y Donaciones)', level=1)
    doc.add_paragraph(
        'Si recibes una herencia o donación, Impuestify calcula el Impuesto de Sucesiones '
        'y Donaciones aplicando la tarifa estatal y las bonificaciones de tu comunidad autónoma. '
        'Más de 8 CCAA tienen bonificaciones del 95-99% para familiares directos.'
    )

    # --- 16. Retenciones ---
    doc.add_heading('16. Retenciones IRPF en facturas', level=1)
    doc.add_paragraph(
        'Si facturas a empresas o profesionales, debes aplicar retención de IRPF en tus facturas. '
        'Impuestify te ayuda a calcular el importe correcto.'
    )

    add_table(doc,
        ['Situación', 'Retención'],
        [
            ['Profesional (general)', '15%'],
            ['Nuevo autónomo (primeros 3 años)', '7%'],
            ['Actividades agrícolas/ganaderas', '2%'],
            ['Actividades forestales', '2%'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ejemplo: ')
    run.bold = True
    p.add_run('"Si facturo 1.000 EUR + IVA con retención del 15%, cuánto cobra mi cliente?"')

    doc.add_page_break()

    # --- 17. Workspaces ---
    doc.add_heading('17. Workspaces de documentos', level=1)
    doc.add_paragraph(
        'Los workspaces son carpetas inteligentes donde puedes subir tus documentos fiscales. '
        'La IA analiza cada workspace de forma aislada, manteniendo el contexto separado.'
    )

    doc.add_heading('Cómo usar los workspaces', level=2)
    add_step(doc, 1, 'Crea un workspace',
             'Ve a "Workspaces" en el menú lateral. Pulsa "Crear workspace" y dale un nombre '
             'descriptivo (ej: "IVA Q1 2026", "Facturas clientes", "Nóminas 2025").')
    add_step(doc, 2, 'Sube documentos',
             'Puedes subir archivos arrastrando y soltando sobre la zona de archivos, o '
             'pulsando el botón "Subir archivo". Formatos aceptados: PDF, Word, Excel, imágenes.')
    add_step(doc, 3, 'Selecciona el workspace en el chat',
             'En la parte superior del chat, selecciona el workspace activo con el desplegable. '
             'Verás un indicador que confirma cuántos documentos se están consultando.')
    add_step(doc, 4, 'Pregunta sobre tus documentos',
             'La IA tendrá el contexto completo de los documentos del workspace seleccionado. '
             'Por ejemplo: "Analiza las facturas de este trimestre y calcula el IVA."')

    doc.add_heading('Ideas de organización', level=2)
    ideas = [
        '"IVA Q1 2026" — Facturas de compras y ventas del primer trimestre',
        '"Nóminas 2025" — Todas las nóminas del año para proyección IRPF',
        '"Notificaciones AEAT" — Notificaciones recibidas de Hacienda',
        '"Gastos deducibles" — Tickets y facturas de gastos profesionales',
    ]
    for idea in ideas:
        doc.add_paragraph(idea, style='List Bullet')

    # --- 17b. Modelos 720 y 721 ---
    doc.add_heading('17b. Modelos 720 y 721 (bienes en el extranjero)', level=1)
    doc.add_paragraph(
        'Si tienes bienes o derechos en el extranjero por valor superior a 50.000 €, '
        'estás obligado a informar a Hacienda mediante el Modelo 720 (declaración informativa). '
        'El Modelo 721 es su equivalente para criptomonedas en el extranjero.'
    )
    doc.add_paragraph('Impuestify te ayuda con:')
    model_720_items = [
        'Determinar si estás obligado a presentar el 720 o el 721',
        'Calcular el valor de tus bienes en el extranjero',
        'Identificar las tres categorías: cuentas bancarias, valores/seguros e inmuebles',
        'Recordarte el plazo de presentación (enero a marzo del año siguiente)',
    ]
    for item in model_720_items:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    run = p.add_run('Importante: ')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 50, 50)
    p.add_run('Desde la sentencia del TJUE de 2022, las sanciones desproporcionadas del 720 '
              'fueron anuladas, pero la obligación de informar sigue vigente.')

    # --- 17c. Clasificador de Facturas ---
    doc.add_heading('17c. Clasificador de Facturas (NUEVO)', level=1)
    doc.add_paragraph(
        'El Clasificador de Facturas permite subir facturas en PDF o foto (JPG, PNG) '
        'y obtener una extracción automática de todos los datos relevantes gracias a '
        'Gemini 3 Flash. Accede desde el menú > "Clasificador Facturas" o directamente '
        'en /clasificador-facturas. Disponible para los planes Autónomo y Creador.'
    )

    doc.add_heading('Cómo funciona', level=2)
    add_step(doc, 1, 'Sube tu factura',
             'Arrastra y suelta un archivo PDF o una foto de la factura en la zona de carga. '
             'También puedes pulsar el botón "Subir factura" para seleccionar el archivo '
             'desde tu dispositivo. Se aceptan formatos PDF, JPG y PNG.')
    add_step(doc, 2, 'Extracción automática con IA',
             'Gemini 3 Flash analiza el documento y extrae automáticamente: '
             'proveedor/emisor, NIF/CIF, fecha de emisión, número de factura, '
             'base imponible, tipo de IVA, cuota de IVA, total factura, '
             'concepto/descripción y forma de pago.')
    add_step(doc, 3, 'Clasificación PGC automática',
             'El sistema asigna automáticamente la cuenta del Plan General Contable (PGC) '
             'correspondiente. Por ejemplo, una factura de material de oficina se clasifica '
             'en la cuenta 629 (Otros servicios) o 602 (Compras de otros aprovisionamientos), '
             'y el IVA soportado en la cuenta 472.')
    add_step(doc, 4, 'Asiento contable generado',
             'Se genera automáticamente el asiento contable (apunte en el Libro Diario) '
             'con las cuentas del Debe y del Haber correspondientes, listo para contabilizar.')

    doc.add_heading('Corregir la clasificación', level=2)
    doc.add_paragraph(
        'Si la cuenta PGC asignada no es correcta, puedes modificarla manualmente. '
        'Pulsa sobre la cuenta sugerida y selecciona la cuenta correcta del desplegable. '
        'El sistema aprende de tus correcciones para mejorar futuras clasificaciones.'
    )

    doc.add_heading('Gestionar facturas', level=2)
    doc.add_paragraph('Desde la vista del clasificador puedes:')
    invoice_actions = [
        'Ver todas las facturas subidas con su estado de clasificación',
        'Filtrar por fecha, proveedor, cuenta PGC o importe',
        'Editar los datos extraídos si la IA no los capturó correctamente',
        'Eliminar facturas que ya no necesites',
        'Exportar la lista de facturas a CSV o Excel',
    ]
    for action in invoice_actions:
        doc.add_paragraph(action, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Consejo: ')
    run.bold = True
    p.add_run(
        'Para mejores resultados, sube facturas con buena resolución y sin rotaciones. '
        'Las fotos tomadas con el móvil funcionan correctamente siempre que el texto sea legible.'
    )

    # --- 17d. Contabilidad y Libros ---
    doc.add_heading('17d. Contabilidad y Libros (NUEVO)', level=1)
    doc.add_paragraph(
        'La sección de Contabilidad te permite consultar tus libros contables generados '
        'automáticamente a partir de las facturas clasificadas. Accede desde el menú > '
        '"Contabilidad" o directamente en /contabilidad. '
        'Disponible para los planes Autónomo y Creador.'
    )

    doc.add_heading('4 libros disponibles', level=2)
    doc.add_paragraph(
        'La vista de contabilidad tiene 4 pestañas, una para cada libro:'
    )

    accounting_books = [
        ('Libro Diario', 'Registro cronológico de todos los asientos contables. '
         'Cada factura clasificada genera un asiento con fecha, número de asiento, '
         'cuentas del Debe y del Haber, importes y concepto. Es el libro principal '
         'y obligatorio para cualquier actividad económica.'),
        ('Libro Mayor', 'Agrupa los movimientos por cuenta contable. Permite ver '
         'todos los apuntes de una cuenta concreta (por ejemplo, todos los gastos '
         'de la cuenta 629 — Otros servicios) con su saldo acumulado. '
         'Ideal para revisar el detalle de cada partida.'),
        ('Balance de Sumas y Saldos', 'Resumen de todas las cuentas con el total '
         'de movimientos en el Debe, total en el Haber y saldo resultante. '
         'Permite verificar que la contabilidad cuadra (suma de Debes = suma de Haberes). '
         'Es la herramienta de control por excelencia.'),
        ('Pérdidas y Ganancias', 'Cuenta de resultados que muestra los ingresos '
         'y gastos del periodo, agrupados por naturaleza según el PGC. '
         'El resultado final indica el beneficio o pérdida de la actividad. '
         'Es el informe que necesitas para el Impuesto de Sociedades y para '
         'el depósito de cuentas en el Registro Mercantil.'),
    ]
    for title, desc in accounting_books:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.color.rgb = RGBColor(26, 86, 219)
        p.add_run(desc)

    doc.add_heading('Exportar datos', level=2)
    doc.add_paragraph(
        'Cada libro puede exportarse en dos formatos:'
    )
    export_formats = [
        'CSV — Formato texto separado por comas, compatible con cualquier programa de hojas de cálculo',
        'Excel (.xlsx) — Formato nativo de Microsoft Excel, con formato y cabeceras incluidas',
    ]
    for fmt in export_formats:
        doc.add_paragraph(fmt, style='List Bullet')
    doc.add_paragraph(
        'Los archivos exportados están preparados para su presentación ante el Registro Mercantil '
        'o para enviar a tu gestor/asesor contable.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Aviso importante: ')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 50, 50)
    p.add_run(
        'La contabilidad generada por Impuestify es una herramienta de apoyo y orientación. '
        'Para el depósito oficial de cuentas anuales en el Registro Mercantil o para '
        'obligaciones formales, recomendamos la revisión por un profesional contable. '
        'Impuestify no sustituye a un software de contabilidad homologado ni al asesoramiento '
        'profesional en materia contable.'
    )

    doc.add_page_break()

    # =========================================================================
    # PARTE V — REFERENCIA
    # =========================================================================
    p = doc.add_paragraph()
    run = p.add_run('PARTE V')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Referencia')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # --- 18. Forales ---
    doc.add_heading('18. Territorios forales', level=1)
    doc.add_paragraph(
        'Impuestify es el ÚNICO asistente fiscal digital que cubre los territorios forales. '
        'Estos territorios tienen un sistema IRPF completamente independiente del régimen común, '
        'con tramos, tipos y deducciones propias.'
    )

    add_table(doc,
        ['Territorio', 'Hacienda', 'Deducciones propias', 'Particularidad'],
        [
            ['Araba/Álava', 'Hacienda Foral de Álava', '15 deducciones',
             'Deducción vivienda vigente (eliminada en régimen común en 2013)'],
            ['Bizkaia', 'Hacienda Foral de Bizkaia', '11 deducciones',
             'Tipos IRPF propios, diferente al estatal'],
            ['Gipuzkoa', 'Hacienda Foral de Gipuzkoa', '11 deducciones',
             'Sistema TicketBAI obligatorio'],
            ['Navarra', 'Hacienda Foral de Navarra', '13 deducciones',
             'IRPF navarro completamente independiente'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Nota: ')
    run.bold = True
    p.add_run('Si resides en un territorio foral, las deducciones estatales NO te aplican. '
              'Impuestify lo gestiona automáticamente: cuando indicas que vives en Araba, Bizkaia, '
              'Gipuzkoa o Navarra, solo te muestra las deducciones forales correspondientes.')

    # --- 19. FAQ ---
    doc.add_heading('19. Preguntas frecuentes', level=1)

    faqs = [
        ('¿Es seguro compartir mis datos fiscales?',
         'Sí. Impuestify cumple con el RGPD y la LSSI-CE. Tus datos se almacenan cifrados, '
         'no compartimos información con terceros, y tenemos filtros automáticos de PII '
         '(datos personales) para evitar que datos sensibles se envíen a la IA.'),
        ('¿Puede la IA equivocarse?',
         'Como cualquier herramienta, la IA puede cometer errores. Por eso citamos siempre '
         'la fuente legal de cada respuesta para que puedas verificarla. Para situaciones '
         'complejas, recomendamos consultar con un asesor fiscal.'),
        ('¿Puedo cancelar mi suscripción?',
         'Sí, en cualquier momento. Ve a Ajustes > Suscripción > "Gestionar suscripción". '
         'Se abre el portal de Stripe donde puedes cancelar sin penalización.'),
        ('¿Funciona en el móvil?',
         'Sí. Impuestify funciona en cualquier dispositivo con navegador moderno: ordenador, '
         'tableta o móvil. Además, es una PWA (Progressive Web App): puedes instalarlo en tu '
         'móvil desde el navegador pulsando "Añadir a pantalla de inicio" y se comportará como '
         'una aplicación nativa.'),
        ('¿Puedo adjuntar documentos en el chat?',
         'Sí. Pulsa el icono de clip junto al campo de texto para adjuntar PDF, imágenes o '
         'documentos. El sistema los analiza automáticamente, extrae los datos y anonimiza la '
         'información personal antes de procesarla.'),
        ('¿Cómo actualizo mi perfil fiscal desde el chat?',
         'Sube un documento (nómina, factura) y pide: "Guarda estos datos en mi perfil fiscal". '
         'La IA extraerá los datos relevantes y los guardará automáticamente.'),
        ('¿Qué pasa con mis datos si cancelo?',
         'Tus datos se conservan 30 días tras la cancelación. Puedes solicitar la eliminación '
         'completa en cualquier momento desde Ajustes > Privacidad.'),
        ('¿La IA puede presentar mi declaración?',
         'Todavía no. Estamos tramitando ser Colaborador Social de la AEAT para poder presentar '
         'declaraciones directamente. Mientras tanto, te preparamos toda la información para que '
         'la presentes tú en 5 minutos vía Renta WEB.'),
        ('¿Puedo compartir una conversación con mi asesor?',
         'Sí. Pulsa el botón de compartir en cualquier conversación y se generará un enlace '
         'público con los datos personales anonimizados. Tu asesor podrá ver la conversación '
         'completa sin acceder a tu cuenta.'),
        ('¿Qué es la calculadora de retenciones?',
         'Es una herramienta pública y gratuita que aplica el algoritmo oficial de la AEAT '
         'para calcular el tipo de retención IRPF que debe aparecer en tu nómina. '
         'No requiere suscripción.'),
        ('¿Cubren los territorios forales?',
         'Sí. Somos el único asistente fiscal digital que cubre los cuatro territorios forales '
         '(Araba, Bizkaia, Gipuzkoa y Navarra) con sus tramos IRPF, deducciones y normativa propia.'),
        ('¿Cómo funciona el clasificador de facturas?',
         'Sube una factura en PDF o foto y Gemini 3 Flash extrae automáticamente todos los datos '
         '(proveedor, NIF, importe, IVA, concepto) y asigna la cuenta del Plan General Contable. '
         'Si la clasificación no es correcta, puedes corregirla manualmente. '
         'Disponible en los planes Autónomo y Creador.'),
        ('¿Los libros contables sirven para el Registro Mercantil?',
         'Los libros generados (Diario, Mayor, Balance, PyG) son una herramienta de apoyo. '
         'Puedes exportarlos a CSV o Excel para enviárselos a tu gestor. '
         'Para el depósito oficial de cuentas anuales, recomendamos la revisión por un profesional contable.'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.color.rgb = RGBColor(26, 86, 219)
        doc.add_paragraph(a)

    # --- 20. Soporte ---
    doc.add_heading('20. Soporte y contacto', level=1)
    doc.add_paragraph('Si necesitas ayuda, tienes varias opciones:')
    support = [
        'Chat de la app — pregunta directamente a la IA, que también puede ayudarte con dudas sobre el funcionamiento',
        'Email — contacto@impuestify.com',
        'Web — impuestify.com/contacto',
    ]
    for s in support:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('impuestify.com')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 219)

    # Save
    output_path = os.path.join(OUTPUT_DIR, "Impuestify_Manual_Usuario_2026.docx")
    doc.save(output_path)
    print(f"User Manual saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_manual()
