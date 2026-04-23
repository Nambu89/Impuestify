"""Generador de la presentacion comercial Impuestify en formato DOCX.

Construye docs/Impuestify_Presentacion_Cliente.docx siguiendo el diseño de
la skill oficial docx (Anthropic): heading styles, tablas con shading,
bullets, header y footer con marca, portada tipo hero.

Uso:
    python scripts/build_presentacion_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


BRAND_BLUE = RGBColor(0x1A, 0x56, 0xDB)
BRAND_CYAN = RGBColor(0x06, 0xB6, 0xD4)
BRAND_DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY_LIGHT = "F1F5F9"
GRAY_ROW = "F8FAFC"
BRAND_HEX = "1A56DB"


def shade_cell(cell, hex_color: str) -> None:
    """Aplica color de fondo a una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None, size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int, *, color: RGBColor = BRAND_BLUE) -> None:
    heading = doc.add_heading(level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(text)
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(13)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def add_bullet(doc: Document, text: str, *, size: int = 11) -> None:
    para = doc.add_paragraph(style="List Bullet")
    run = para.runs[0] if para.runs else para.add_run(text)
    if not para.runs:
        para.add_run(text)
    for run in para.runs:
        run.font.size = Pt(size)
    para.runs[-1].text = text


def add_bullet_simple(doc: Document, text: str, size: int = 11) -> None:
    para = doc.add_paragraph(text, style="List Bullet")
    for run in para.runs:
        run.font.size = Pt(size)


def add_feature_row(doc: Document, features: list[tuple[str, str]]) -> None:
    """Añade una tabla 3xN con features: cada feature = (titulo, descripcion)."""
    # Reorganizar en filas de 3 columnas
    rows_data = [features[i:i + 3] for i in range(0, len(features), 3)]
    for row_data in rows_data:
        while len(row_data) < 3:
            row_data.append(("", ""))
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, (title, desc) in enumerate(row_data):
            cell = table.rows[0].cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if title:
                shade_cell(cell, GRAY_LIGHT)
                cell.text = ""
                p1 = cell.paragraphs[0]
                run_title = p1.add_run(title)
                run_title.bold = True
                run_title.font.size = Pt(11)
                run_title.font.color.rgb = BRAND_BLUE
                p2 = cell.add_paragraph()
                run_desc = p2.add_run(desc)
                run_desc.font.size = Pt(9)
        set_column_widths(table, [Cm(5.7), Cm(5.7), Cm(5.7)])
        doc.add_paragraph()  # espacio entre filas


def set_column_widths(table, widths: list) -> None:
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = widths[idx]


def add_table_with_header(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list | None = None, highlight_rows: list[int] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        shade_cell(cell, BRAND_HEX)
        set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)

    highlight_rows = highlight_rows or []
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        is_highlight = r_idx in highlight_rows
        shade = GRAY_ROW if (r_idx % 2 == 1 and not is_highlight) else None
        if is_highlight:
            shade = "FEF3C7"
        for c_idx, value in enumerate(row):
            cell = tr.cells[c_idx]
            if shade:
                shade_cell(cell, shade)
            set_cell_text(cell, value, bold=is_highlight, size=9)

    if col_widths:
        set_column_widths(table, col_widths)


def add_horizontal_rule(doc: Document) -> None:
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), BRAND_HEX)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Impuestify — Presentación comercial")
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_BLUE
    run.bold = True

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Documento confidencial · impuestify.com · © 2026 Impuestify")
    run.font.size = Pt(8)
    run.font.color.rgb = BRAND_DARK


def build_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand = p_brand.add_run("IMPUESTIFY")
    brand.bold = True
    brand.font.size = Pt(48)
    brand.font.color.rgb = BRAND_BLUE

    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag = p_tag.add_run("El copiloto fiscal con IA para toda España")
    tag.italic = True
    tag.font.size = Pt(16)
    tag.font.color.rgb = BRAND_CYAN

    for _ in range(2):
        doc.add_paragraph()

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = p_sub.add_run("PRESENTACIÓN COMERCIAL")
    sub.bold = True
    sub.font.size = Pt(14)
    sub.font.color.rgb = BRAND_DARK

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date = p_date.add_run("Cliente potencial · Abril 2026")
    date.font.size = Pt(11)
    date.font.color.rgb = BRAND_DARK

    for _ in range(10):
        doc.add_paragraph()

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p_footer.add_run("impuestify.com")
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.color.rgb = BRAND_BLUE

    doc.add_page_break()


def build_section_1(doc: Document) -> None:
    add_heading(doc, "1. Qué es Impuestify", level=1)
    add_paragraph(
        doc,
        "Impuestify es una plataforma fiscal con inteligencia artificial para España que "
        "cubre de forma integral las obligaciones tributarias de particulares, autónomos, "
        "creadores digitales y sociedades.",
    )
    add_paragraph(
        doc,
        "Desde un único espacio web accesible en impuestify.com, el usuario puede:",
    )
    for item in [
        "Calcular su IRPF con total precisión territorial.",
        "Clasificar y contabilizar sus facturas automáticamente.",
        "Generar borradores de modelos oficiales.",
        "Preparar su defensa frente a requerimientos y liquidaciones de la Administración.",
        "Gestionar su contabilidad completa para el Registro Mercantil.",
        "Resolver dudas fiscales conversacionalmente con respuestas basadas en legislación oficial verificable.",
    ]:
        add_bullet_simple(doc, item)


def build_section_2(doc: Document) -> None:
    add_heading(doc, "2. El problema que resolvemos", level=1)
    add_paragraph(
        doc,
        "La fiscalidad española es uno de los sistemas tributarios más fragmentados de Europa:",
    )
    for item in [
        "21 regímenes fiscales distintos (15 CCAA + 4 forales + Ceuta/Melilla).",
        "Más de 1.000 deducciones activas entre estatales, autonómicas y forales.",
        "Normativa actualizada cada año por AEAT, BOE y Diputaciones Forales.",
        "Procedimientos administrativos complejos que requieren asesoramiento especializado.",
    ]:
        add_bullet_simple(doc, item)
    add_paragraph(
        doc,
        "La oferta actual se divide entre gestorías tradicionales caras y herramientas "
        "digitales de cobertura parcial (normalmente solo régimen común). Impuestify es la "
        "primera solución que combina IA avanzada con cobertura fiscal total del territorio "
        "español.",
    )


def build_section_3(doc: Document) -> None:
    add_heading(doc, "3. Herramientas incluidas", level=1)

    features = [
        ("Chat fiscal inteligente", "Asistente conversacional que responde en lenguaje natural, citando la legislación oficial que justifica cada respuesta. Streaming en tiempo real."),
        ("Simulador IRPF completo", "Calcula el resultado de la renta contemplando trabajo, ahorro, inmuebles, actividades económicas, alquileres, criptomonedas, pérdidas y tributación conjunta."),
        ("Guía Fiscal adaptativa", "Asistente paso a paso que se adapta al perfil del usuario (Particular, Autónomo o Creador) y a su Comunidad Autónoma. Estimador en tiempo real."),
        ("DefensIA — Defensor fiscal", "Motor de defensa tributaria con anti-alucinación. Analiza requerimientos y liquidaciones, genera borradores de escritos (reposición, TEAR, alegaciones)."),
        ("Clasificador de facturas", "OCR automático de facturas (PDF, JPG, HEIC) con clasificación contable al Plan General Contable y generación de asientos en partida doble."),
        ("Contabilidad PGC completa", "Libro Diario, Libro Mayor, Balance de Situación y Cuenta de Pérdidas y Ganancias automáticos. Exportación CSV / Excel para Registro Mercantil."),
        ("Workspaces contables", "Espacios de trabajo con dashboard visual: KPIs, evolución trimestral de IVA, top proveedores, gráficos de ingresos y gastos, tablas PGC."),
        ("Modelo 200 — Sociedades", "Simulador completo del IS para SL, SA y empresas de nueva creación. Siete territorios. Pagos fraccionados Modelo 202. Borrador PDF."),
        ("Generador de modelos oficiales", "Borradores PDF de modelos 303, 130, 308, 720, 721, IPSI y variantes forales (300 Gipuzkoa, F69 Navarra, 420 IGIC Canarias)."),
        ("Seis calculadoras públicas", "Acceso gratuito sin registro: sueldo neto, retenciones IRPF, umbrales contables, modelos obligatorios, obligado a declarar, checklist del borrador."),
        ("Calendario fiscal", "Recordatorios automatizados de vencimientos de modelos trimestrales y anuales, y plazos de pago."),
        ("Análisis de nóminas y AEAT", "Subida de nóminas y notificaciones AEAT en PDF con interpretación automática de conceptos, retenciones, bases imponibles y plazos."),
    ]
    add_feature_row(doc, features)


def build_section_4(doc: Document) -> None:
    add_heading(doc, "4. Cobertura fiscal — 21 territorios al 100 %", level=1)

    headers = ["Territorio", "Régimen IRPF", "Modelo IVA", "Sucesiones/Donaciones"]
    rows = [
        ["Andalucía", "Común", "303 / 390", "Sí"],
        ["Aragón", "Común", "303 / 390", "Sí"],
        ["Asturias", "Común", "303 / 390", "Sí"],
        ["Baleares", "Común", "303 / 390", "Sí"],
        ["Canarias", "Común", "420 (IGIC)", "Sí"],
        ["Cantabria", "Común", "303 / 390", "Sí"],
        ["Castilla-La Mancha", "Común", "303 / 390", "Sí"],
        ["Castilla y León", "Común", "303 / 390", "Sí"],
        ["Cataluña", "Común", "303 / 390", "Sí"],
        ["Extremadura", "Común", "303 / 390", "Sí"],
        ["Galicia", "Común", "303 / 390", "Sí"],
        ["La Rioja", "Común", "303 / 390", "Sí"],
        ["Madrid", "Común", "303 / 390", "Sí"],
        ["Murcia", "Común", "303 / 390", "Sí"],
        ["Valencia", "Común", "303 / 390", "Sí"],
        ["Álava", "Foral (normativa propia)", "303", "Normativa foral"],
        ["Bizkaia", "Foral (normativa propia)", "303", "Normativa foral"],
        ["Gipuzkoa", "Foral (7 tramos)", "300", "Normativa foral"],
        ["Navarra", "Foral (11 tramos)", "F69", "Normativa foral"],
        ["Ceuta", "Bonificación 60 %", "IPSI (6 tipos)", "Específica"],
        ["Melilla", "Bonificación 60 %", "IPSI (6 tipos)", "Específica"],
    ]
    highlight = list(range(15, 21))
    add_table_with_header(doc, headers, rows, col_widths=[Cm(4.5), Cm(5), Cm(3.5), Cm(4)], highlight_rows=highlight)

    add_paragraph(
        doc,
        "Somos el único sistema del mercado con cobertura íntegra de los territorios forales "
        "y de Ceuta/Melilla.",
        bold=True,
    )


def build_section_5(doc: Document) -> None:
    add_heading(doc, "5. Base documental indexada", level=1)
    add_paragraph(
        doc,
        "Impuestify opera sobre una base de legislación oficial indexada de gran volumen que "
        "se mantiene actualizada de forma automática:",
    )
    headers = ["Métrica", "Valor"]
    rows = [
        ["Documentos oficiales indexados", "463"],
        ["Fragmentos de texto procesados", "92.393"],
        ["Vectores semánticos generados", "85.587"],
        ["Fuentes oficiales integradas", "AEAT, BOE, Diputaciones Forales, Boletines autonómicos"],
        ["URLs monitorizadas por el crawler", "90"],
        ["Territorios con documentación específica", "23"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(7), Cm(10)])
    add_paragraph(
        doc,
        "Los contenidos se actualizan automáticamente cuando cambia la normativa oficial. "
        "Cada respuesta del asistente se acompaña de citas verificables a la fuente "
        "legislativa correspondiente.",
    )


def build_section_6(doc: Document) -> None:
    add_heading(doc, "6. Seguridad y privacidad", level=1)
    add_paragraph(
        doc,
        "La plataforma está diseñada con un modelo de defensa en profundidad de 13 capas que "
        "protege tanto los datos personales del usuario como la integridad de las respuestas "
        "fiscales:",
    )
    headers = ["Capa", "Función"]
    rows = [
        ["Autenticación", "JWT + refresh tokens + MFA (TOTP) + CAPTCHA"],
        ["Protección frente a manipulación", "Filtros de prompt injection y moderación de contenido"],
        ["Detección automática de datos personales", "DNI, NIE, IBAN, teléfonos, correos"],
        ["Prevención de inyección SQL", "Consultas parametrizadas + detección OWASP"],
        ["Rate limiting", "Protección frente a abuso y denegación de servicio"],
        ["Cumplimiento GDPR", "Derechos de borrado, rectificación y portabilidad"],
        ["Auditoría", "Registro inmutable de accesos y operaciones sensibles"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(6.5), Cm(10.5)])
    add_paragraph(
        doc,
        "Cumplimiento normativo: RGPD, AEPD, LSSI-CE. Datos alojados en infraestructura cloud europea.",
        bold=True,
    )


def build_section_7(doc: Document) -> None:
    add_heading(doc, "7. Calidad del producto", level=1)
    headers = ["Métrica", "Valor"]
    rows = [
        ["Tests automatizados en backend", "~1.800"],
        ["Tests automatizados en frontend", "Suite completa Vitest"],
        ["Tests del motor de defensa DefensIA", "379"],
        ["Tests del módulo Modelo 200 IS", "47"],
        ["Tests end-to-end de usuario", "Playwright multi-dispositivo"],
        ["Tiempo de respuesta del simulador IRPF", "< 100 ms (sin IA)"],
        ["Disponibilidad del servicio", "99,9 %"],
        ["Despliegue", "Automático en cada actualización"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(8), Cm(9)])


def build_section_8(doc: Document) -> None:
    add_heading(doc, "8. Modelo 200 — Impuesto sobre Sociedades", level=1)

    add_paragraph(
        doc,
        "Impuestify incorpora un módulo dedicado al Impuesto sobre Sociedades (Modelo 200) "
        "diseñado para sociedades limitadas (SL), sociedades anónimas (SA) y empresas de "
        "nueva creación. Es el único simulador del mercado español que integra en una sola "
        "herramienta el régimen común, los cuatro territorios forales, la Zona Especial "
        "Canaria y Ceuta/Melilla.",
    )

    add_heading(doc, "Cobertura territorial del módulo IS", level=2)

    headers_t = ["Territorio", "Normativa aplicada", "Tipo general", "Régimen especial"]
    rows_t = [
        ["Régimen común (15 CCAA)", "LIS + modificaciones estatales", "25 %", "Empresa nueva creación 15 %"],
        ["Álava", "Norma Foral propia", "24 %", "Reducida dimensión 20 %"],
        ["Bizkaia", "Norma Foral propia", "24 %", "Reducida dimensión 20 %"],
        ["Gipuzkoa", "Norma Foral propia", "24 %", "Reducida dimensión 20 %"],
        ["Navarra", "Ley Foral propia", "28 %", "Reducida dimensión 23 %"],
        ["Zona Especial Canaria (ZEC)", "Régimen fiscal especial", "4 %", "Ventaja competitiva brutal"],
        ["Ceuta / Melilla", "Bonificación estatal", "12,5 %", "Bonificación 50 % sobre tipo"],
    ]
    add_table_with_header(doc, headers_t, rows_t, col_widths=[Cm(4.5), Cm(5), Cm(3), Cm(4.5)])

    add_heading(doc, "Funcionalidades del módulo", level=2)
    for item in [
        "Simulador completo del Modelo 200 con cálculo de base imponible, ajustes "
        "extracontables, compensación de bases negativas y cuota líquida.",
        "Pagos fraccionados del ejercicio mediante Modelo 202 (Art. 40 LIS) en sus tres "
        "modalidades, con cálculo automático de la opción óptima.",
        "Generación de borrador PDF con las 16 casillas principales del Modelo 200 listas "
        "para revisión antes de la presentación telemática.",
        "Integración bidireccional con los workspaces contables: la cuenta de Pérdidas y "
        "Ganancias del ejercicio se vuelca automáticamente en el simulador.",
        "Aplicación automática de tipos reducidos: empresas de reducida dimensión, entidades "
        "de nueva creación, entidades sin fines lucrativos y entidades ZEC.",
        "Wizard interactivo de 4 pasos que guía al usuario desde la selección de territorio "
        "hasta la descarga del borrador.",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Datos precargables desde contabilidad", level=2)
    add_paragraph(
        doc,
        "Si la sociedad utiliza los workspaces contables de Impuestify durante el ejercicio, "
        "al llegar el momento de calcular el Modelo 200 los siguientes campos se rellenan "
        "automáticamente:",
    )
    for item in [
        "Resultado contable del ejercicio (grupos 6 y 7 del PGC).",
        "Ingresos de explotación y gastos deducibles agrupados por epígrafe.",
        "Amortizaciones y deterioros aplicados a inmovilizado.",
        "Bases imponibles negativas pendientes de compensar de ejercicios anteriores.",
        "Retenciones e ingresos a cuenta soportados durante el ejercicio.",
    ]:
        add_bullet_simple(doc, item)

    add_heading(doc, "Valor diferencial frente a competencia", level=2)

    headers_d = ["Característica", "Impuestify", "Software IS estándar", "Gestoría tradicional"]
    rows_d = [
        ["Régimen común + 4 forales + ZEC + Ceuta/Melilla", "Sí", "No", "Parcial"],
        ["Integración con contabilidad del ejercicio", "Sí (automática)", "No", "Manual"],
        ["Cálculo automático de opción óptima del Modelo 202", "Sí", "Parcial", "Manual"],
        ["Aplicación de tipos reducidos y bonificaciones", "Sí (automática)", "Manual", "Manual"],
        ["PDF borrador del Modelo 200 en 4 pasos", "Sí", "No", "No"],
        ["Tests automatizados del motor de cálculo", "47", "—", "—"],
    ]
    add_table_with_header(doc, headers_d, rows_d, col_widths=[Cm(5.5), Cm(3.5), Cm(4), Cm(4)])

    add_heading(doc, "Casos de uso del módulo", level=2)

    cases = [
        ("SL de nueva creación en Madrid",
         "Una SL con primer ejercicio cerrado selecciona Madrid en el wizard. Impuestify aplica "
         "automáticamente el tipo reducido del 15 % para entidades de nueva creación y genera "
         "el borrador del Modelo 200 con la bonificación correctamente aplicada."),
        ("SL de reducida dimensión en Bizkaia",
         "Una SL con facturación inferior a 10 millones de euros en Bizkaia obtiene el tipo "
         "foral reducido del 20 % aplicado sobre la base imponible foral, con el tratamiento "
         "específico de la Diputación Foral en los ajustes extracontables."),
        ("Entidad acogida a Zona Especial Canaria",
         "Una empresa ZEC que cumple los requisitos de actividad, empleo y volumen de inversión "
         "aplica el tipo reducido del 4 %. El simulador valida el cumplimiento de las "
         "condiciones y calcula el ahorro fiscal frente al régimen común."),
    ]
    for title, body in cases:
        add_heading(doc, title, level=3)
        add_paragraph(doc, body)


def build_section_9(doc: Document) -> None:
    add_heading(doc, "9. Casos de uso reales", level=1)

    cases = [
        ("Particular — IRPF con múltiples pagadores",
         "Mariana, asalariada en Madrid con ingresos de dos empresas, quiere saber si está "
         "obligada a declarar. Introduce los pagadores en la guía fiscal, Impuestify aplica "
         "los umbrales del Art. 96 LIRPF y devuelve obligación con la cuota estimada a pagar."),
        ("Autónomo — Gestión contable trimestral",
         "Iván, diseñador freelance en Gipuzkoa, sube sus 40 facturas trimestrales en PDF. "
         "Impuestify las clasifica en el PGC, genera los asientos, calcula el Modelo 300 "
         "(equivalente al 303 en Gipuzkoa) y exporta el Libro Diario para su asesor."),
        ("Creador — IVA por plataforma",
         "Lucía, streamer con ingresos de YouTube, Twitch y Patreon, necesita declarar IVA "
         "intracomunitario (Modelo 349) por los ingresos de plataformas irlandesas. Impuestify "
         "identifica el origen, aplica la regla y genera el borrador."),
        ("Pyme — Impuesto sobre Sociedades",
         "Una SL de Bizkaia quiere estimar su Impuesto sobre Sociedades del ejercicio. El "
         "simulador aplica el tipo foral correcto, contempla las bonificaciones aplicables y "
         "genera el borrador del Modelo 200 precargado con los datos contables del workspace."),
        ("Defensa fiscal — Requerimiento AEAT",
         "Carlos recibe una propuesta de liquidación paralela de la AEAT por rendimientos de "
         "alquiler. Sube el PDF a DefensIA, el sistema extrae los datos, identifica la fase "
         "procesal, verifica la documentación aplicable y genera un borrador de escrito de "
         "reposición en lenguaje jurídico formal."),
    ]
    for title, body in cases:
        add_heading(doc, title, level=3)
        add_paragraph(doc, body)


def build_section_10(doc: Document) -> None:
    add_heading(doc, "10. Por qué Impuestify", level=1)
    headers = ["Característica", "Impuestify", "Gestoría tradicional", "Software fiscal estándar"]
    rows = [
        ["Cobertura 21 territorios fiscales", "Sí", "Parcial (según firma)", "No (habitualmente régimen común)"],
        ["Territorios forales (País Vasco, Navarra)", "Sí", "Parcial", "No"],
        ["Ceuta y Melilla", "Sí", "Parcial", "No"],
        ["IA conversacional con citas verificables", "Sí", "No", "Parcial"],
        ["OCR de facturas automatizado", "Sí", "No", "Parcial"],
        ["Motor de defensa fiscal con escritos", "Sí", "Parcial (servicio adicional)", "No"],
        ["Cobertura Modelo 200 (IS) multi-territorial", "Sí", "Parcial (común)", "Sí (manual)"],
        ["Acceso 24/7", "Sí", "No", "Sí"],
        ["Actualización automática de normativa", "Sí", "Parcial", "Parcial"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(5.5), Cm(3.5), Cm(4), Cm(4)])


def build_section_11(doc: Document) -> None:
    add_heading(doc, "11. Disponibilidad y acceso", level=1)
    for item in [
        "Web: impuestify.com",
        "Plataforma: navegador web (PC, Mac, móvil iOS, Android).",
        "Instalación PWA: sí, con notificaciones push para recordatorios fiscales.",
        "Idioma: español (cobertura completa). Inglés en roadmap 2026.",
        "Registro: correo electrónico + contraseña.",
        "Prueba: las seis calculadoras públicas no requieren registro.",
    ]:
        add_bullet_simple(doc, item)


def build_section_12(doc: Document) -> None:
    add_heading(doc, "12. Aviso legal", level=1)
    add_paragraph(
        doc,
        "Impuestify es una herramienta de asistencia informativa apoyada en legislación "
        "oficial vigente.",
    )
    for item in [
        "No constituye asesoramiento fiscal profesional. Las respuestas y cálculos son orientativos. "
        "Cada situación personal puede tener matices que requieran atención profesional.",
        "El módulo DefensIA genera borradores orientativos que deben ser revisados y validados por "
        "un abogado o asesor fiscal colegiado antes de su presentación ante la administración.",
        "Cumplimiento GDPR: los datos del usuario son tratados conforme al Reglamento General de "
        "Protección de Datos. El usuario mantiene los derechos de acceso, rectificación, supresión, "
        "portabilidad y oposición.",
    ]:
        add_bullet_simple(doc, item)


def build_section_13(doc: Document) -> None:
    add_heading(doc, "13. Contacto", level=1)

    add_paragraph(doc, "Fernando Prada", bold=True, size=13)
    add_paragraph(doc, "Fundador", italic=True, size=11)

    for label, value in [
        ("Web", "impuestify.com"),
        ("Correo", "fernando.prada@proton.me"),
        ("LinkedIn", "(a completar)"),
        ("Soporte", "formulario de contacto en la propia plataforma"),
    ]:
        para = doc.add_paragraph()
        r_label = para.add_run(f"{label}:  ")
        r_label.bold = True
        r_label.font.size = Pt(11)
        r_value = para.add_run(value)
        r_value.font.size = Pt(11)

    add_horizontal_rule(doc)
    add_paragraph(doc, "Documento confidencial. Para evaluación comercial.", italic=True, size=9)
    add_paragraph(doc, "© 2026 Impuestify — Todos los derechos reservados.", italic=True, size=9)


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    set_header_footer(doc)

    build_cover(doc)
    build_section_1(doc)
    build_section_2(doc)
    doc.add_page_break()
    build_section_3(doc)
    doc.add_page_break()
    build_section_4(doc)
    doc.add_page_break()
    build_section_5(doc)
    build_section_6(doc)
    doc.add_page_break()
    build_section_7(doc)
    doc.add_page_break()
    build_section_8(doc)
    doc.add_page_break()
    build_section_9(doc)
    doc.add_page_break()
    build_section_10(doc)
    build_section_11(doc)
    build_section_12(doc)
    doc.add_page_break()
    build_section_13(doc)

    out_path = Path(__file__).resolve().parent.parent / "docs" / "Impuestify_Presentacion_Cliente.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"DOCX generado en: {out_path}")
    print(f"Tamano: {out_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
