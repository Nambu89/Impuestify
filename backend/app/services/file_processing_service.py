"""
File Processing Service for TaxIA

Handles file uploads, classification, text extraction, and embeddings for Workspaces.
Supports: PDF, DOCX, Excel, CSV, Images (JPG/PNG).
Integrates specialized extractors for invoices and payslips.
"""
import csv
import io
import json
import logging
import os
import uuid
from dataclasses import asdict
from datetime import datetime, timezone
from typing import Optional, Dict, Any

from fastapi import UploadFile

from app.database.turso_client import get_db_client
from app.security.document_integrity import document_integrity_scanner
from app.utils.pdf_extractor import (
    extract_pdf_text,
    extract_pdf_text_plain,
    extract_image_text,
    PDFExtractionResult,
)

logger = logging.getLogger(__name__)


class FileProcessingService:
    """Service for processing uploaded workspace files."""

    ACCEPTED_TYPES = {
        "application/pdf": "pdf",
        "image/jpeg": "image",
        "image/png": "image",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "excel",
        "application/vnd.ms-excel": "excel",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/csv": "csv",
        "application/csv": "csv",
    }

    # Whether to generate embeddings (can be disabled for faster uploads)
    ENABLE_EMBEDDINGS = True

    async def process_file_upload(
        self,
        workspace_id: str,
        file: UploadFile,
        file_type_hint: Optional[str] = None,
        user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process an uploaded file: save metadata, extract content, and generate embeddings.

        Args:
            workspace_id: ID of the target workspace
            file: FastAPI UploadFile object
            file_type_hint: Optional hint for file type (nomina, factura, etc.)

        Returns:
            Dictionary with created file metadata
        """
        if file.content_type not in self.ACCEPTED_TYPES:
            raise ValueError(f"File type {file.content_type} not supported")

        file_id = str(uuid.uuid4())

        # Read content
        content = await file.read()
        file_size = len(content)

        # Classify file type
        file_category = file_type_hint or self._classify_file_type(file.filename)

        # Extract text and structured data
        extracted_text = ""
        extracted_data = {}
        processing_status = "pending"

        try:
            mime = file.content_type or ""
            fmt = self.ACCEPTED_TYPES.get(mime, "")

            if mime == "application/pdf":
                # Use plain text for payslips/invoices (tabular PDFs) —
                # pymupdf4llm markdown mangles column layout making regex impossible.
                # Use markdown for other docs (better structure for RAG).
                if file_category in ("nomina", "factura"):
                    result: PDFExtractionResult = await extract_pdf_text_plain(content, file.filename)
                else:
                    result: PDFExtractionResult = await extract_pdf_text(content, file.filename)

                if result.success:
                    extracted_text = result.markdown_text
                    processing_status = "completed"

                    # Apply specialized extractors based on file type
                    if file_category == "factura":
                        extracted_data = await self._extract_invoice_data(extracted_text)
                    elif file_category == "nomina":
                        extracted_data = await self._extract_payslip_data(extracted_text)
                    elif file_category == "declaracion":
                        extracted_data = self._extract_declaration_data(extracted_text)
                    else:
                        extracted_data = {"pages": result.total_pages}
                else:
                    processing_status = "error"
                    logger.error(f"PDF extraction failed: {result.error}")

            elif fmt == "docx":
                extracted_text = self._extract_docx_text(content, file.filename)
                processing_status = "completed" if extracted_text else "error"

            elif fmt == "excel":
                extracted_text = self._extract_excel_text(content, file.filename)
                processing_status = "completed" if extracted_text else "error"

            elif fmt == "csv":
                extracted_text = self._extract_csv_text(content, file.filename)
                processing_status = "completed" if extracted_text else "error"

            elif fmt == "image":
                img_result = await extract_image_text(content, file.filename)
                if img_result.success:
                    extracted_text = img_result.markdown_text
                    processing_status = "completed"
                else:
                    processing_status = "error"
                    logger.error(f"Image extraction failed: {img_result.error}")

            # Guard: empty extraction should not be marked as completed
            if processing_status == "completed" and (not extracted_text or len(extracted_text.strip()) < 10):
                processing_status = "error"
                extracted_data = {"error": "No se pudo extraer texto del documento"}
                logger.warning(f"Empty extraction for {file.filename}, setting status=error")

        except Exception as e:
            logger.error(f"Error extracting text from {file.filename}: {e}")
            processing_status = "error"

        # Document Integrity Scan (Capa 13) — fail open: scanner errors never block processing
        integrity_score: Optional[float] = None
        integrity_findings: Optional[str] = None

        if extracted_text:
            try:
                scan_result = document_integrity_scanner.scan(extracted_text, source="upload")

                # Merge metadata findings if the PDF validator produced metadata
                pdf_metadata = getattr(file, "_integrity_metadata", None)
                if pdf_metadata and isinstance(pdf_metadata, dict):
                    meta_findings = document_integrity_scanner.scan_metadata(pdf_metadata)
                    if meta_findings:
                        scan_result.findings.extend(meta_findings)
                        from app.security.document_integrity import _compute_risk_score
                        scan_result.risk_score = round(_compute_risk_score(scan_result.findings), 4)
                        scan_result.is_safe = scan_result.risk_score < 0.3

                integrity_score = scan_result.risk_score
                integrity_findings = json.dumps([asdict(f) for f in scan_result.findings]) if scan_result.findings else None

                if scan_result.risk_score > 0.8:
                    # BLOCK: do not pass to agents, do not generate embeddings
                    logger.error(
                        "Document integrity BLOCK: file=%s risk_score=%.2f findings=%d — "
                        "processing_status set to blocked_integrity",
                        file.filename, scan_result.risk_score, len(scan_result.findings),
                    )
                    processing_status = "blocked_integrity"

                elif scan_result.risk_score >= 0.6:
                    # SANITIZE: replace critical/high fragments before passing to agents
                    logger.warning(
                        "Document integrity SANITIZE: file=%s risk_score=%.2f — "
                        "sanitizing extracted text",
                        file.filename, scan_result.risk_score,
                    )
                    extracted_text = document_integrity_scanner.sanitize(extracted_text, scan_result.findings)

                elif scan_result.risk_score >= 0.3:
                    # WARN: process normally but log
                    logger.warning(
                        "Document integrity WARN: file=%s risk_score=%.2f findings=%d",
                        file.filename, scan_result.risk_score, len(scan_result.findings),
                    )
                # else: PASS — process normally, no log needed

            except Exception as scan_exc:
                # Fail open: log and continue without integrity data
                logger.error(
                    "Document integrity scan failed for %s (fail-open): %s",
                    file.filename, scan_exc,
                )

        # Save to Database
        db = await get_db_client()
        now = datetime.now(timezone.utc).isoformat()

        await db.execute(
            """
            INSERT INTO workspace_files (
                id, workspace_id, filename, file_type, mime_type,
                file_size, extracted_text, extracted_data, processing_status,
                integrity_score, integrity_findings, created_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            [
                file_id,
                workspace_id,
                file.filename,
                file_category,
                file.content_type,
                file_size,
                extracted_text,
                json.dumps(extracted_data) if isinstance(extracted_data, dict) and extracted_data else None,
                processing_status,
                integrity_score,
                integrity_findings,
                now
            ]
        )

        # Generate embeddings asynchronously (if enabled and extraction succeeded)
        # Blocked documents must NOT generate embeddings
        embedding_status = "skipped"
        if self.ENABLE_EMBEDDINGS and processing_status == "completed" and extracted_text:
            try:
                embedding_result = await self._generate_file_embeddings(
                    db, workspace_id, file_id, extracted_text, file.filename
                )
                embedding_status = "completed" if embedding_result.get('success') else "failed"
            except Exception as e:
                logger.error(f"Embedding generation failed for {file.filename}: {e}")
                embedding_status = "error"

        # Auto-classify invoice into PGC after processing
        if (
            file_category == "factura"
            and extracted_data
            and not extracted_data.get("error")
            and user_id
            and processing_status == "completed"
        ):
            try:
                await self._auto_classify_invoice(file_id, user_id, extracted_data, db)
            except Exception as e:
                logger.warning(f"Auto-classification failed for file {file_id}: {e}")
                # Don't fail the upload — classification is best-effort

        return {
            "id": file_id,
            "filename": file.filename,
            "file_type": file_category,
            "status": processing_status,
            "size": file_size,
            "extracted_data": extracted_data,
            "embedding_status": embedding_status,
            "integrity_score": integrity_score,
            "integrity_findings": integrity_findings,
        }

    def _classify_file_type(self, filename: str) -> str:
        """Classify file type from filename."""
        lower_name = filename.lower()

        # Payslip patterns
        if any(p in lower_name for p in ["nomina", "nómina", "payslip", "salario"]):
            return "nomina"

        # Invoice patterns
        if any(p in lower_name for p in ["factura", "invoice", "fra", "fact"]):
            return "factura"

        # Tax declaration patterns
        if any(p in lower_name for p in [
            "declaracion", "declaración", "modelo",
            "303", "130", "420", "300",
            "390", "100", "200",
        ]):
            return "declaracion"

        return "otro"

    async def _extract_invoice_data(self, text: str) -> Dict[str, Any]:
        """Extract structured data from invoice text."""
        try:
            from app.services.invoice_extractor import get_invoice_extractor

            extractor = get_invoice_extractor()
            data = await extractor.extract_from_text(text)

            # Add summary
            data['summary'] = extractor.generate_summary(data)
            data['vat_breakdown'] = extractor.get_vat_breakdown(data)

            return data

        except Exception as e:
            logger.error(f"Invoice extraction failed: {e}")
            return {"error": str(e)}

    async def _extract_payslip_data(self, text: str) -> Dict[str, Any]:
        """Extract structured data from payslip text."""
        try:
            from app.services.payslip_extractor import PayslipExtractor

            extractor = PayslipExtractor()
            data = extractor._parse_payslip_data(text)

            # Add summary
            data['summary'] = extractor.generate_summary(data)
            data['stats'] = extractor.get_extraction_stats(data)

            return data

        except Exception as e:
            logger.error(f"Payslip extraction failed: {e}")
            return {"error": str(e)}

    # ------------------------------------------------------------------
    # Multi-format text extraction
    # ------------------------------------------------------------------

    def _extract_docx_text(self, content: bytes, filename: str) -> str:
        """Extract text from a DOCX file (paragraphs + tables)."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            parts: list[str] = []

            # Paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append(" | ".join(cells))

            result = "\n".join(parts)
            logger.info(f"DOCX extracted {len(result)} chars from {filename}")
            return result

        except Exception as e:
            logger.error(f"DOCX extraction failed for {filename}: {e}")
            return ""

    def _extract_excel_text(self, content: bytes, filename: str) -> str:
        """Extract text from an Excel file (all sheets, all cells)."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            parts: list[str] = []

            for sheet in wb.worksheets:
                parts.append(f"=== Hoja: {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    # Skip completely empty rows
                    if any(cells):
                        parts.append(" | ".join(cells))

            wb.close()
            result = "\n".join(parts)
            logger.info(f"Excel extracted {len(result)} chars from {filename}")
            return result

        except Exception as e:
            logger.error(f"Excel extraction failed for {filename}: {e}")
            return ""

    def _extract_csv_text(self, content: bytes, filename: str) -> str:
        """Extract text from a CSV file."""
        try:
            # Try UTF-8 first, then Latin-1 (common in Spanish docs)
            for encoding in ("utf-8", "latin-1", "cp1252"):
                try:
                    text = content.decode(encoding)
                    break
                except (UnicodeDecodeError, ValueError):
                    continue
            else:
                text = content.decode("utf-8", errors="replace")

            reader = csv.reader(io.StringIO(text))
            parts: list[str] = []
            for row in reader:
                if any(cell.strip() for cell in row):
                    parts.append(" | ".join(row))

            result = "\n".join(parts)
            logger.info(f"CSV extracted {len(result)} chars from {filename}")
            return result

        except Exception as e:
            logger.error(f"CSV extraction failed for {filename}: {e}")
            return ""

    def _extract_declaration_data(self, text: str) -> Dict[str, Any]:
        """Extract structured data from tax declaration text (303/130/420)."""
        try:
            from app.services.declaration_extractor import DeclarationExtractor

            extractor = DeclarationExtractor()
            result = extractor.extract(text)

            if not result.success:
                return {"error": result.error, "type": "declaracion"}

            return {
                "type": "declaracion",
                "modelo": result.modelo,
                "metadata": result.metadata,
                "fields": result.fields,
                "territory": result.territory,
                "confidence": result.confidence,
                "casillas_count": len(result.casillas_raw),
            }

        except Exception as e:
            logger.error(f"Declaration extraction failed: {e}")
            return {"error": str(e), "type": "declaracion"}

    async def _auto_classify_invoice(
        self,
        workspace_file_id: str,
        user_id: str,
        extracted_data: Dict[str, Any],
        db,
    ) -> None:
        """
        Auto-classify a workspace invoice into a PGC account and create
        a libro_registro entry with clasificacion_confianza='pendiente_confirmacion'.
        """
        from app.config import settings
        from app.services.invoice_classifier_service import InvoiceClassifierService
        from app.services.contabilidad_service import ContabilidadService

        if not settings.GOOGLE_GEMINI_API_KEY:
            logger.warning("Gemini API key not configured — skipping auto-classification")
            return

        # Extract fields from the structured extracted_data
        concepto = extracted_data.get("summary") or extracted_data.get("concepto") or ""
        emisor_nombre = extracted_data.get("emisor_nombre") or extracted_data.get("emisor", {}).get("nombre", "")
        base_imponible = float(extracted_data.get("base_imponible") or extracted_data.get("base_imponible_total") or 0)
        tipo_iva = float(extracted_data.get("tipo_iva") or extracted_data.get("tipo_iva_pct") or 0)
        cuota_iva = float(extracted_data.get("cuota_iva") or 0)
        total = float(extracted_data.get("total") or base_imponible)
        numero_factura = extracted_data.get("numero_factura") or ""
        fecha_factura = extracted_data.get("fecha_factura") or datetime.now(timezone.utc).strftime("%Y-%m-%d")
        emisor_nif = extracted_data.get("emisor_nif") or extracted_data.get("emisor", {}).get("nif_cif", "")
        receptor_nif = extracted_data.get("receptor_nif") or extracted_data.get("receptor", {}).get("nif_cif", "")
        receptor_nombre = extracted_data.get("receptor_nombre") or extracted_data.get("receptor", {}).get("nombre", "")
        retencion_irpf = float(extracted_data.get("retencion_irpf") or 0)
        retencion_irpf_pct = float(extracted_data.get("retencion_irpf_pct") or 0)
        tipo_re = float(extracted_data.get("tipo_re_pct") or extracted_data.get("tipo_re") or 0)
        cuota_re = float(extracted_data.get("cuota_re") or 0)

        if base_imponible <= 0:
            logger.info("Skipping auto-classification: base_imponible is 0 for file %s", workspace_file_id)
            return

        # Default tipo: recibida (expense) — most common for uploaded invoices
        tipo = extracted_data.get("tipo") or "recibida"

        # PGC classification via Gemini
        classifier = InvoiceClassifierService(
            api_key=settings.GOOGLE_GEMINI_API_KEY,
            db=db,
            model=settings.GEMINI_MODEL,
        )

        clasificacion = await classifier.classify(
            concepto=concepto,
            emisor_nombre=emisor_nombre,
            tipo=tipo,
            base_imponible=base_imponible,
        )

        # Parse date for year / trimestre
        try:
            fecha_dt = datetime.strptime(fecha_factura, "%Y-%m-%d")
        except ValueError:
            try:
                fecha_dt = datetime.strptime(fecha_factura, "%d/%m/%Y")
            except ValueError:
                fecha_dt = datetime.now(timezone.utc)

        year = fecha_dt.year
        trimestre = (fecha_dt.month - 1) // 3 + 1

        # Insert into libro_registro with pendiente_confirmacion
        invoice_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        raw_extraction_json = json.dumps(extracted_data, default=str)

        await db.execute(
            """
            INSERT INTO libro_registro
                (id, user_id, workspace_file_id, tipo, numero_factura,
                 fecha_factura, emisor_nif, emisor_nombre,
                 receptor_nif, receptor_nombre, concepto,
                 base_imponible, tipo_iva, cuota_iva,
                 tipo_re, cuota_re, retencion_irpf_pct, retencion_irpf,
                 total, cuenta_pgc, cuenta_pgc_nombre,
                 clasificacion_confianza, raw_extraction,
                 year, trimestre, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            [
                invoice_id,
                user_id,
                workspace_file_id,
                tipo,
                numero_factura,
                fecha_factura,
                emisor_nif,
                emisor_nombre,
                receptor_nif,
                receptor_nombre,
                concepto[:200] if concepto else "",
                base_imponible,
                tipo_iva,
                cuota_iva,
                tipo_re,
                cuota_re,
                retencion_irpf_pct,
                retencion_irpf,
                total,
                clasificacion.cuenta_code,
                clasificacion.cuenta_nombre,
                "pendiente_confirmacion",
                raw_extraction_json,
                year,
                trimestre,
                now,
            ],
        )

        # Generate + save asiento contable
        concepto_asiento = f"Factura {numero_factura}" if numero_factura else concepto[:80]
        asiento_lines = ContabilidadService.generate_asiento_lines(
            tipo=tipo,
            cuenta_pgc_code=clasificacion.cuenta_code,
            cuenta_pgc_nombre=clasificacion.cuenta_nombre,
            base_imponible=base_imponible,
            cuota_iva=cuota_iva,
            total=total,
            retencion_irpf=retencion_irpf,
            concepto=concepto_asiento,
        )

        contabilidad = ContabilidadService(db=db)
        await contabilidad.save_asiento(
            user_id=user_id,
            libro_registro_id=invoice_id,
            fecha=fecha_factura,
            lines=asiento_lines,
            year=year,
            trimestre=trimestre,
        )

        logger.info(
            "Auto-classified workspace file %s → cuenta %s (%s) for user %s",
            workspace_file_id, clasificacion.cuenta_code, clasificacion.cuenta_nombre, user_id,
        )

    async def _generate_file_embeddings(
        self,
        db,
        workspace_id: str,
        file_id: str,
        text: str,
        filename: str
    ) -> Dict[str, Any]:
        """Generate embeddings for file content."""
        try:
            from app.services.workspace_embedding_service import get_workspace_embedding_service

            service = get_workspace_embedding_service()
            result = await service.embed_workspace_file(
                db, workspace_id, file_id, text, filename
            )

            return result

        except ImportError:
            logger.warning("Embedding service not available")
            return {"success": False, "error": "Service not available"}
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            return {"success": False, "error": str(e)}

    async def reprocess_file(self, file_id: str, user_id: str = None) -> Dict[str, Any]:
        """
        Reprocess an existing file (re-extract and re-embed).

        Args:
            file_id: ID of the file to reprocess
            user_id: ID of the requesting user (ownership verification)

        Returns:
            Updated file metadata
        """
        db = await get_db_client()

        if user_id:
            # Verify ownership: file must belong to a workspace owned by this user
            ownership_result = await db.execute(
                """
                SELECT wf.id FROM workspace_files wf
                JOIN workspaces w ON wf.workspace_id = w.id
                WHERE wf.id = ? AND w.user_id = ?
                """,
                [file_id, user_id]
            )
            if not ownership_result.rows:
                raise ValueError("File not found or access denied")

        # Get file info
        result = await db.execute(
            "SELECT * FROM workspace_files WHERE id = ?",
            [file_id]
        )

        if not result.rows:
            raise ValueError("File not found or access denied")

        file_info = result.rows[0]

        # If we have extracted text, re-run specialized extractors
        extracted_text = file_info.get('extracted_text', '')
        file_category = file_info.get('file_type', 'otro')
        workspace_id = file_info['workspace_id']

        extracted_data = {}

        if extracted_text:
            if file_category == "factura":
                extracted_data = await self._extract_invoice_data(extracted_text)
            elif file_category == "nomina":
                extracted_data = await self._extract_payslip_data(extracted_text)
            elif file_category == "declaracion":
                extracted_data = self._extract_declaration_data(extracted_text)

        # Update database
        await db.execute(
            """
            UPDATE workspace_files
            SET extracted_data = ?, processing_status = 'completed'
            WHERE id = ?
            """,
            [json.dumps(extracted_data) if isinstance(extracted_data, dict) and extracted_data else None, file_id]
        )

        # Regenerate embeddings
        if self.ENABLE_EMBEDDINGS and extracted_text:
            # Delete old embeddings
            from app.services.workspace_embedding_service import get_workspace_embedding_service
            service = get_workspace_embedding_service()
            await service.delete_file_embeddings(db, file_id)

            # Generate new
            await self._generate_file_embeddings(
                db, workspace_id, file_id, extracted_text, file_info['filename']
            )

        return {
            "id": file_id,
            "status": "reprocessed",
            "extracted_data": extracted_data
        }


# Global instance
file_processing_service = FileProcessingService()
