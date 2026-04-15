"""DefensIA Export Service (T2B-007 + T2B-008).

Convierte escritos markdown en DOCX (python-docx) y PDF (reportlab + DejaVuSans).
Incluye disclaimer obligatorio en header y footer del documento exportado
(superficies 3 y 4 de las 4 del disclaimer DefensIA).

Pin exacto: reportlab + DejaVuSans TTFont. NO weasyprint (L4 del plan-checker).
Si la fuente DejaVuSans no está disponible, fallback a Helvetica builtin
(soporta Latin-1, cubre las tildes castellanas y la eñe).
"""
from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


DISCLAIMER_CANONICO = (
    "DefensIA es una herramienta de asistencia tecnica que no constituye "
    "asesoramiento juridico vinculante. Revisa y adapta el contenido antes "
    "de presentarlo ante cualquier administracion."
)


class DefensiaExportService:
    """Export markdown → DOCX / PDF con disclaimer integrado en header + footer."""

    # ------------------------------------------------------------------
    # T2B-007: DOCX export
    # ------------------------------------------------------------------

    def markdown_a_docx(
        self,
        markdown_text: str,
        *,
        titulo: str = "Escrito DefensIA",
        disclaimer: str = DISCLAIMER_CANONICO,
        autor: str = "Impuestify DefensIA",
    ) -> bytes:
        """Convierte markdown a DOCX bytes con disclaimer en header y footer.

        Args:
            markdown_text: Contenido del escrito en markdown básico (headings #/##/###,
                bullets con `- `, párrafos separados por línea en blanco).
            titulo: Título del documento (metadata `core_properties.title`).
            disclaimer: Texto del disclaimer a insertar en header y footer.
            autor: Autor del documento (metadata `core_properties.author`).

        Returns:
            Bytes del fichero DOCX listos para descarga o almacenamiento.
        """
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as e:
            raise RuntimeError(
                "python-docx no disponible. Instala python-docx>=1.1.2"
            ) from e

        doc = Document()

        # Metadata del documento
        doc.core_properties.author = autor
        doc.core_properties.title = titulo

        # Header con disclaimer
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = disclaimer
        for run in header_para.runs:
            run.font.size = Pt(8)
            run.font.italic = True

        # Footer con disclaimer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = disclaimer
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.italic = True

        # Cuerpo markdown básico: headings + párrafos + bullets
        for linea in markdown_text.split("\n"):
            linea_stripped = linea.strip()
            if linea_stripped.startswith("# "):
                doc.add_heading(linea_stripped[2:], level=1)
            elif linea_stripped.startswith("## "):
                doc.add_heading(linea_stripped[3:], level=2)
            elif linea_stripped.startswith("### "):
                doc.add_heading(linea_stripped[4:], level=3)
            elif linea_stripped.startswith("- "):
                doc.add_paragraph(linea_stripped[2:], style="List Bullet")
            elif linea_stripped:
                doc.add_paragraph(linea_stripped)
            else:
                # Línea en blanco → párrafo vacío para mantener separación
                doc.add_paragraph("")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # ------------------------------------------------------------------
    # T2B-008: PDF export
    # ------------------------------------------------------------------

    def markdown_a_pdf(
        self,
        markdown_text: str,
        *,
        titulo: str = "Escrito DefensIA",
        disclaimer: str = DISCLAIMER_CANONICO,
        autor: str = "Impuestify DefensIA",
    ) -> bytes:
        """Convierte markdown a PDF bytes con DejaVuSans (o Helvetica fallback).

        Usa reportlab + registro manual de TTFont DejaVuSans para soportar tildes
        y eñe. Si DejaVuSans no está disponible en el sistema, cae a Helvetica
        builtin (soporta Latin-1 y cubre las tildes castellanas).

        Args:
            markdown_text: Contenido del escrito en markdown básico.
            titulo: Título del PDF (metadata del documento).
            disclaimer: Texto del disclaimer a insertar al inicio y al final.
            autor: Autor del PDF (metadata).

        Returns:
            Bytes del fichero PDF listos para descarga.
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as e:
            raise RuntimeError(
                "reportlab no disponible. Instala reportlab>=4.0"
            ) from e

        # Registrar fuente Unicode (DejaVuSans o fallback Helvetica)
        font_name = self._registrar_fuente_unicode()

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=25 * mm,
            rightMargin=25 * mm,
            topMargin=25 * mm,
            bottomMargin=25 * mm,
            title=titulo,
            author=autor,
        )

        styles = getSampleStyleSheet()
        style_body = ParagraphStyle(
            "DefensiaBody",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=14,
        )
        style_h1 = ParagraphStyle(
            "DefensiaH1",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceAfter=8,
        )
        style_h2 = ParagraphStyle(
            "DefensiaH2",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=12,
            leading=16,
            spaceAfter=6,
        )
        style_h3 = ParagraphStyle(
            "DefensiaH3",
            parent=styles["Heading3"],
            fontName=font_name,
            fontSize=11,
            leading=14,
            spaceAfter=4,
        )
        style_disclaimer = ParagraphStyle(
            "DefensiaDisclaimer",
            parent=styles["Italic"],
            fontName=font_name,
            fontSize=8,
            leading=10,
        )

        story: list = []

        # Disclaimer inicial (superficie 3)
        story.append(Paragraph(self._escape_xml(disclaimer), style_disclaimer))
        story.append(Spacer(1, 4 * mm))

        # Cuerpo del markdown
        for linea in markdown_text.split("\n"):
            linea_stripped = linea.strip()
            if linea_stripped.startswith("# "):
                story.append(
                    Paragraph(self._escape_xml(linea_stripped[2:]), style_h1)
                )
            elif linea_stripped.startswith("## "):
                story.append(
                    Paragraph(self._escape_xml(linea_stripped[3:]), style_h2)
                )
            elif linea_stripped.startswith("### "):
                story.append(
                    Paragraph(self._escape_xml(linea_stripped[4:]), style_h3)
                )
            elif linea_stripped.startswith("- "):
                story.append(
                    Paragraph(
                        "• " + self._escape_xml(linea_stripped[2:]),
                        style_body,
                    )
                )
            elif linea_stripped:
                story.append(
                    Paragraph(self._escape_xml(linea_stripped), style_body)
                )
            else:
                story.append(Spacer(1, 3 * mm))

        # Disclaimer final (superficie 4)
        story.append(Spacer(1, 6 * mm))
        story.append(Paragraph(self._escape_xml(disclaimer), style_disclaimer))

        doc.build(story)
        return buffer.getvalue()

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _escape_xml(text: str) -> str:
        """Escape de caracteres XML especiales para reportlab Paragraph."""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    def _registrar_fuente_unicode(self) -> str:
        """Registra DejaVuSans si está disponible, sino fallback a Helvetica.

        Búsqueda en rutas comunes de Linux, macOS y Windows. Si no se encuentra,
        devuelve 'Helvetica' (fuente builtin de reportlab, soporta Latin-1
        cubriendo tildes castellanas y la eñe).

        Returns:
            Nombre de la fuente registrada ("DejaVuSans" o "Helvetica").
        """
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            logger.warning(
                "reportlab pdfmetrics no disponible, usando Helvetica builtin"
            )
            return "Helvetica"

        # Si ya está registrada (ej. en una llamada previa), devolver directo
        try:
            registered = pdfmetrics.getRegisteredFontNames()
            if "DejaVuSans" in registered:
                return "DejaVuSans"
        except Exception:
            pass

        candidates = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/TTF/DejaVuSans.ttf",
            "/Library/Fonts/DejaVuSans.ttf",
            "C:\\Windows\\Fonts\\DejaVuSans.ttf",
        ]
        for path in candidates:
            try:
                if Path(path).exists():
                    pdfmetrics.registerFont(TTFont("DejaVuSans", path))
                    return "DejaVuSans"
            except Exception as e:
                logger.warning(
                    "Error registrando DejaVuSans desde %s: %s", path, e
                )
                continue

        logger.info(
            "DejaVuSans no encontrado en el sistema, usando Helvetica builtin "
            "(soporta tildes Latin-1)"
        )
        return "Helvetica"
