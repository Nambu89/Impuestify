"""Generate Word document with marketing plan answers for Erika."""
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

title = doc.add_heading('Respuestas para el Plan de Marketing de Impuestify', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Preparado por: Fernando Prada')
doc.add_paragraph('Fecha: 27 de marzo de 2026')
doc.add_paragraph('Para: Erika Cepeda')
doc.add_paragraph('')

# ========== 1 ==========
doc.add_heading('1. Que quiero conseguir ahora mismo con la web', level=1)
doc.add_paragraph('Orden de prioridad (de mas a menos importante):')

for t, d in [
    ('1. Que la gente pruebe la herramienta',
     'Es lo mas importante. Impuestify se vende sola cuando la pruebas. El asistente fiscal con IA, la guia fiscal interactiva y el simulador IRPF son el mejor argumento de venta.'),
    ('2. Que se registren usuarios',
     'Necesito masa critica de usuarios para validar el producto y generar traccion. El registro es gratuito y da acceso inmediato.'),
    ('3. Conseguir confianza y credibilidad',
     'Al ser un producto nuevo sobre un tema sensible (impuestos), la confianza es clave. Tenemos cumplimiento RGPD, AI Act, LSSI-CE y LOPDGDD visible en toda la web.'),
    ('4. Que se hable de Impuestify',
     'El boca a boca y la visibilidad organica son fundamentales en esta etapa. Un usuario satisfecho que lo recomiende vale mas que cualquier anuncio.'),
    ('5. Que entre mucha gente en la web',
     'El trafico importa, pero prefiero 100 personas que prueben la herramienta a 10.000 que solo lean el landing.'),
    ('6. Empezar a captar clientes',
     'Es importante pero no urgente. Primero quiero que la gente conozca el producto y confie. La conversion a pago vendra cuando el valor este demostrado.'),
]:
    p = doc.add_paragraph()
    p.add_run(t).bold = True
    p.add_run(f'\n{d}')

# ========== 2 ==========
doc.add_heading('2. A quien quiero ayudar primero', level=1)
doc.add_paragraph('Orden de prioridad:')

for t, d in [
    ('1. Gente que tiene dudas con la renta',
     'Es el publico mas amplio y con necesidad mas inmediata (campana de renta empieza el 8 de abril de 2026). Son millones de personas que cada ano se enfrentan a la declaracion con miedo a equivocarse.'),
    ('2. Particulares',
     'Asalariados, pensionistas, gente con hipoteca, hijos, deducciones... Son el grueso de los declarantes y nuestro plan de 5 EUR/mes esta pensado para ellos.'),
    ('3. Autonomos',
     'Necesitan ayuda constante (IVA trimestral, IRPF, cuotas SS, modelos 303/130). Plan de 39 EUR/mes. Son menos en volumen pero con mayor valor por cliente.'),
    ('4. Personas perdidas con impuestos o IVA',
     'Gente que no sabe que modelo presentar, cuando, o cuanto va a pagar. Impuestify les da respuestas inmediatas.'),
    ('5. Creadores de contenido',
     'YouTubers, streamers, influencers. Nicho especifico con necesidades complejas (IVA intracomunitario, Modelo 349, DAC7). Plan Creator de 49 EUR/mes. Segmento en crecimiento pero mas pequeno.'),
]:
    p = doc.add_paragraph()
    p.add_run(t).bold = True
    p.add_run(f'\n{d}')

# ========== 3 ==========
doc.add_heading('3. Que problema resuelve Impuestify mejor que nada', level=1)

for problem in [
    'Simplifica temas fiscales que normalmente son un caos. La fiscalidad espanola tiene 21 territorios diferentes (17 CCAA + 4 forales + Ceuta/Melilla), cada uno con sus propias deducciones, escalas y modelos. Impuestify es el unico asistente que cubre los 21.',
    'Ayuda a no perder deducciones. Tenemos mas de 1.000 deducciones en nuestra base de datos (estatales + autonomicas + forales). El sistema las descubre automaticamente segun tu perfil fiscal y tu CCAA.',
    'Aclara que hay que presentar y cuando. El calendario fiscal integrado con alertas push avisa de plazos (modelo 303, 130, 100, 714...). Nunca mas te pierdes un plazo.',
    'Ayuda a entender impuestos sin volverse loco. El asistente fiscal con IA responde en lenguaje natural, con datos concretos y referencias legales. No jerga, no rodeos.',
    'Evita errores. El simulador IRPF calcula tu resultado con precision (escalas estatales + autonomicas + forales), para que sepas ANTES de presentar si te sale a pagar o a devolver.',
]:
    doc.add_paragraph(problem, style='List Bullet')

# ========== 4 ==========
doc.add_heading('4. Que tiene de especial frente a otras opciones', level=1)

doc.add_heading('Frente a buscar en Google:', level=2)
doc.add_paragraph('Google te da resultados genericos, muchas veces desactualizados o de otros paises. Impuestify tiene 454 documentos oficiales indexados (AEAT, BOE, haciendas forales) con 89.000+ fragmentos de legislacion vigente. Las respuestas estan respaldadas por fuentes concretas, no por articulos de blogs.')

doc.add_heading('Frente a preguntarle a ChatGPT:', level=2)
doc.add_paragraph('ChatGPT no conoce tu CCAA, tu situacion laboral ni tus deducciones especificas. Inventa datos fiscales ("alucina"). Impuestify usa RAG (Retrieval-Augmented Generation) con documentacion oficial verificada, conoce los 21 territorios espanoles, y tiene herramientas de calculo real (simulador IRPF, calculadora de cuotas, comparador conjunta/individual). No inventa: calcula.')

doc.add_heading('Frente a usar otra web (TaxDown, Declarando...):', level=2)
doc.add_paragraph('Las webs tipo TaxDown o Declarando son formularios rigidos que te hacen la declaracion por ti (y cobran por ello). Impuestify es un asistente inteligente que te explica, te orienta y te ayuda a entender. No compite con ellos: los complementa. Ademas, Impuestify cubre territorios forales (Pais Vasco, Navarra) que la mayoria de competidores ignoran.')

doc.add_heading('Frente a ir a una gestoria:', level=2)
doc.add_paragraph('Una gestoria cuesta 60-200 EUR por declaracion y solo te atiende en horario laboral. Impuestify esta disponible 24/7, responde en segundos, y cuesta 5 EUR/mes para particulares. No sustituye al asesor fiscal para casos complejos, pero resuelve el 80% de las dudas que la gente tiene antes de necesitar uno.')

# ========== 5 ==========
doc.add_heading('5. Que quiero que haga una persona cuando entra en la web', level=1)
doc.add_paragraph('Orden de prioridad:')

for t, d in [
    ('1. Hacer una simulacion',
     'La guia fiscal interactiva (7 pasos) es la joya de la corona. En 2 minutos sabe si le sale a pagar o a devolver, con deducciones autonomicas incluidas.'),
    ('2. Probar una calculadora',
     'El simulador IRPF, la calculadora de salario neto para autonomos, el comparador conjunta/individual... Son herramientas que enganchan.'),
    ('3. Registrarse',
     'El registro desbloquea el chat con IA, el workspace de documentos y las alertas fiscales. Es el paso a la conversion.'),
    ('4. Contratar algo',
     'Los planes de pago (5/39/49 EUR/mes) son el objetivo final, pero no el primer paso.'),
    ('5. Leer',
     'El contenido informativo (landing, paginas de territorios) atrae trafico SEO pero la accion real es probar la herramienta.'),
    ('6. Dejar su email',
     'No tenemos newsletter por ahora. El registro ya captura el email.'),
]:
    p = doc.add_paragraph()
    p.add_run(t).bold = True
    p.add_run(f'\n{d}')

# ========== 6 ==========
doc.add_heading('6. Que es lo mas potente que ya tenemos', level=1)
doc.add_paragraph('Sin duda, el chat con asistente fiscal IA. Le preguntas cualquier cosa sobre impuestos en Espana y te responde con datos concretos, cifras exactas y referencias legales, personalizado para tu CCAA y tu situacion.')
doc.add_paragraph('Ejemplo real: le preguntas "que epigrafe IAE necesito para ponedoras de huevos en Bizkaia" y te responde con el codigo exacto (019.3 - Avicultura de puesta), con la tabla de cuotas forales y la referencia al BOB. Ningun otro servicio hace esto.')
doc.add_paragraph('Lo segundo mas potente es la guia fiscal interactiva: en 7 pasos tienes tu resultado IRPF estimado con deducciones autonomicas. Visual, rapido, sin tecnicismos.')

# ========== 7 ==========
doc.add_heading('7. Que materiales tenemos ahora mismo', level=1)

for mat, desc in [
    ('Logo:', 'Si. Escudo con cerebro IA + texto "Impuestify". Version horizontal para header, favicon para navegador.'),
    ('Colores e imagen de marca:', 'Si. Azul corporativo (#1e40af), cyan acento (#06b6d4), fondo blanco. Consistente en toda la web.'),
    ('Textos de la web:', 'Si. Landing completa, paginas de territorios (forales, Ceuta/Melilla, Canarias), pagina de creadores, politica privacidad, terminos, transparencia IA.'),
    ('Capturas:', 'Si. Se pueden generar facilmente de todas las secciones.'),
    ('Videos:', 'No. Pendiente. Podria hacer screencasts (grabacion de pantalla mostrando la herramienta).'),
    ('Perfiles en redes:', 'Si. LinkedIn + Instagram + TikTok creados. Contenido minimo por ahora.'),
    ('Base de datos:', 'Si. Usuarios registrados (beta testers activos), 454 documentos fiscales oficiales, 1.000+ deducciones.'),
    ('Contactos:', 'Si. 3 beta testers activos dando feedback real.'),
    ('Testimonios:', 'En proceso. Los beta testers estan usando la herramienta pero aun no tenemos testimonios formales.'),
    ('Presupuesto para publicidad:', 'Limitado. Estoy bootstrapping. Prefiero invertir en contenido organico y SEO antes que en ads pagados.'),
]:
    p = doc.add_paragraph()
    p.add_run(mat).bold = True
    p.add_run(f' {desc}')

# ========== 8 ==========
doc.add_heading('8. Cuanto tiempo real puedo dedicar a esto', level=1)

for t, d in [
    ('Tiempo semanal:', 'Unas 5-8 horas semanales para marketing/contenido. El resto del tiempo lo dedico a desarrollo del producto.'),
    ('Crear contenido:', 'Si, puedo escribir posts y crear graficos basicos. Tengo conocimiento profundo del tema fiscal que es la mayor ventaja.'),
    ('Grabar videos:', 'Prefiero evitarlo por ahora. No me siento comodo delante de la camara. Si puedo hacer screencasts (grabacion de pantalla mostrando la herramienta).'),
    ('Web vs redes sociales:', 'Prefiero que el plan dependa mas de la web y del SEO. Las redes sociales como complemento, no como canal principal. El contenido de la web dura, el de redes desaparece.'),
    ('Que no quiero hacer:', 'No quiero hacer spam, no quiero cold outreach agresivo, no quiero publicar 3 veces al dia en redes. Prefiero calidad a cantidad.'),
]:
    p = doc.add_paragraph()
    p.add_run(t).bold = True
    p.add_run(f' {d}')

# ========== 9 ==========
doc.add_heading('9. Que tono quiero que tenga la marca', level=1)
doc.add_paragraph('Profesional pero sencillo. Experto y claro.')
doc.add_paragraph('Impuestify tutea al usuario. Usa lenguaje natural, evita jerga fiscal innecesaria. Cuando usa terminos tecnicos los explica. Es directo: dato primero, explicacion despues.')
doc.add_paragraph('No es informal ni coloquial (los impuestos son cosa seria), pero tampoco es frio ni burocratico. Es como hablar con un amigo que sabe mucho de fiscalidad y te lo explica sin hacerte sentir tonto.')
doc.add_paragraph('Referencia de tono: como un asesor fiscal joven que te explica las cosas en un cafe, con cifras y referencias legales pero sin abrumarte.')

# ========== 10 ==========
doc.add_heading('10. Que me daria miedo comunicar mal', level=1)

for t, d in [
    ('Prometer cosas que no se puedan asegurar:',
     'Impuestify da calculos orientativos, NO declaraciones oficiales. Siempre incluimos el aviso "Calculo orientativo - consulta con un asesor para tu caso concreto". Nunca debemos decir "te hacemos la renta" o "te garantizamos la devolucion".'),
    ('Dar sensacion de que sustituye a un asesor fiscal:',
     'No lo sustituye. Lo complementa. Impuestify resuelve el 80% de las dudas comunes, pero para herencias complejas, reestructuraciones societarias o inspecciones, necesitas un profesional. Esto hay que comunicarlo siempre.'),
    ('Parecer poco serio:',
     'Los impuestos son dinero real de la gente. No podemos ser frivolos ni usar humor que reste credibilidad. El tono debe transmitir que detras hay un producto robusto con fuentes oficiales.'),
    ('Sonar demasiado agresivo:',
     'No quiero el tipico marketing de "ahorra 3.000 EUR con este truco". Prefiero "descubre las deducciones que te corresponden segun tu CCAA". Informar, no manipular.'),
    ('Prometer cobertura fiscal completa:',
     'Cubrimos los 21 territorios de Espana, pero no todos los supuestos posibles. Hay que ser honestos con las limitaciones.'),
]:
    p = doc.add_paragraph()
    p.add_run(t).bold = True
    p.add_run(f' {d}')

# ========== Datos adicionales ==========
doc.add_heading('Datos adicionales que pueden ser utiles para el plan', level=1)

doc.add_heading('Numeros actuales:', level=2)
for e in [
    'Web: impuestify.com (dominio propio, SSL, hosting en Railway)',
    'Documentacion RAG: 454 documentos oficiales, 89.000+ chunks indexados',
    'Deducciones: 1.000+ en base de datos (21 territorios al 100%)',
    'Tests: 1.212 tests automatizados pasando',
    '13 capas de seguridad (RGPD, AI Act, LlamaGuard, PII detection...)',
    'PWA instalable en movil',
    'Planes: Particular 5 EUR/mes, Autonomo 39 EUR/mes, Creator 49 EUR/mes',
    'Campana de renta 2026: empieza el 8 de abril',
]:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('Fechas clave proximas:', level=2)
for d in [
    '8 abril 2026: Inicio campana de Renta 2025 (oportunidad ENORME de marketing)',
    '30 junio 2026: Fin de campana de Renta',
    '20 abril 2026: Primer trimestre IVA (autonomos)',
    '20 julio 2026: Segundo trimestre IVA',
]:
    doc.add_paragraph(d, style='List Bullet')

# Save
import os
output = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'Impuestify_Plan_Marketing_Respuestas.docx')
doc.save(output)
print(f'Word guardado en: {output}')
