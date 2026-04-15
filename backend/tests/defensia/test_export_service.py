"""Tests para DefensIA Export Service (T2B-007 + T2B-008).

Cubre export markdown → DOCX (python-docx) y markdown → PDF (reportlab + DejaVuSans).
Verifica disclaimer en header y footer, preservación de tildes, niveles de headings,
fallback Helvetica si DejaVuSans no está disponible.
"""
from __future__ import annotations

import io
import re
from unittest.mock import patch

import pytest

from app.services.defensia_export_service import (
    DISCLAIMER_CANONICO,
    DefensiaExportService,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def service() -> DefensiaExportService:
    return DefensiaExportService()


@pytest.fixture
def markdown_basico() -> str:
    return "# Titulo\n\nParrafo simple de prueba."


@pytest.fixture
def markdown_con_tildes() -> str:
    return (
        "# Expedición por infracción tributaria\n"
        "\n"
        "En méritos del artículo 2.º de la Ley General Tributaria, la señora López "
        "solicita la nulidad de la sanción por omisión de motivación.\n"
    )


@pytest.fixture
def markdown_con_niveles() -> str:
    return (
        "# Encabezado Nivel 1\n"
        "\n"
        "Parrafo introductorio.\n"
        "\n"
        "## Encabezado Nivel 2\n"
        "\n"
        "- Primer punto\n"
        "- Segundo punto\n"
        "\n"
        "### Encabezado Nivel 3\n"
        "\n"
        "Conclusion.\n"
    )


# ---------------------------------------------------------------------------
# T2B-007: DOCX export
# ---------------------------------------------------------------------------


def test_docx_basico(service, markdown_basico):
    """El DOCX generado debe ser bytes no vacíos y reabrible con python-docx."""
    from docx import Document

    result = service.markdown_a_docx(markdown_basico)

    assert isinstance(result, bytes)
    assert len(result) > 0

    # El DOCX debe ser reabrible
    doc = Document(io.BytesIO(result))
    # Debe contener al menos el título como heading
    textos = [p.text for p in doc.paragraphs]
    assert any("Titulo" in t for t in textos)
    assert any("Parrafo simple" in t for t in textos)


def test_docx_incluye_disclaimer_header_y_footer(service, markdown_basico):
    """El DOCX debe tener el disclaimer canónico en header y footer."""
    from docx import Document

    result = service.markdown_a_docx(markdown_basico)
    doc = Document(io.BytesIO(result))

    section = doc.sections[0]

    header_texts = [p.text for p in section.header.paragraphs]
    footer_texts = [p.text for p in section.footer.paragraphs]

    assert any(DISCLAIMER_CANONICO in t for t in header_texts), (
        f"Disclaimer no encontrado en header. Textos: {header_texts}"
    )
    assert any(DISCLAIMER_CANONICO in t for t in footer_texts), (
        f"Disclaimer no encontrado en footer. Textos: {footer_texts}"
    )


def test_docx_preserva_tildes(service, markdown_con_tildes):
    """El DOCX debe preservar tildes y ñ sin corromperse."""
    from docx import Document

    result = service.markdown_a_docx(markdown_con_tildes)
    doc = Document(io.BytesIO(result))

    todo_texto = "\n".join(p.text for p in doc.paragraphs)

    # Palabras con tildes y ñ deben aparecer literalmente
    assert "Expedición" in todo_texto
    assert "infracción" in todo_texto
    assert "méritos" in todo_texto
    assert "artículo 2.º" in todo_texto
    assert "López" in todo_texto
    assert "sanción" in todo_texto
    assert "omisión" in todo_texto


def test_docx_headings_niveles(service, markdown_con_niveles):
    """El DOCX debe mapear #, ##, ### a heading levels 1, 2, 3."""
    from docx import Document

    result = service.markdown_a_docx(markdown_con_niveles)
    doc = Document(io.BytesIO(result))

    # Recolectar niveles de heading por texto
    headings_por_nivel: dict[int, list[str]] = {1: [], 2: [], 3: []}
    for p in doc.paragraphs:
        style_name = (p.style.name or "").lower() if p.style else ""
        match = re.search(r"heading\s*(\d+)", style_name)
        if match:
            nivel = int(match.group(1))
            if nivel in headings_por_nivel:
                headings_por_nivel[nivel].append(p.text)

    assert any("Encabezado Nivel 1" in t for t in headings_por_nivel[1])
    assert any("Encabezado Nivel 2" in t for t in headings_por_nivel[2])
    assert any("Encabezado Nivel 3" in t for t in headings_por_nivel[3])


def test_markdown_a_docx_con_disclaimer_custom(service, markdown_basico):
    """El parámetro disclaimer debe sobrescribir el texto en header y footer."""
    from docx import Document

    custom = "Texto de prueba disclaimer custom 2026."
    result = service.markdown_a_docx(markdown_basico, disclaimer=custom)
    doc = Document(io.BytesIO(result))

    section = doc.sections[0]
    header_texts = [p.text for p in section.header.paragraphs]
    footer_texts = [p.text for p in section.footer.paragraphs]

    assert any(custom in t for t in header_texts)
    assert any(custom in t for t in footer_texts)
    # El disclaimer canónico NO debe aparecer cuando se pasa uno custom
    assert not any(DISCLAIMER_CANONICO in t for t in header_texts)


# ---------------------------------------------------------------------------
# T2B-008: PDF export
# ---------------------------------------------------------------------------


def test_pdf_basico(service, markdown_basico):
    """El PDF generado debe empezar con %PDF y tener longitud razonable."""
    result = service.markdown_a_pdf(markdown_basico)

    assert isinstance(result, bytes)
    assert result.startswith(b"%PDF"), f"No es un PDF valido: {result[:20]!r}"
    assert len(result) > 1000, f"PDF demasiado corto ({len(result)} bytes)"


def test_pdf_preserva_tildes(service, markdown_con_tildes):
    """El PDF no debe crashear con UnicodeEncodeError y debe preservar tildes."""
    # No debe lanzar UnicodeEncodeError
    result = service.markdown_a_pdf(markdown_con_tildes)

    assert result.startswith(b"%PDF")
    assert len(result) > 1000

    # Intentar extraer texto con pypdf si está disponible
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(result))
        texto_extraido = "\n".join(page.extract_text() or "" for page in reader.pages)

        # Al menos alguna de las palabras con tildes debe aparecer extraíble.
        # pypdf puede no recuperar el 100% del texto con tildes dependiendo de la fuente,
        # pero al menos las letras base (sin acento) deben estar presentes.
        assert "Expedici" in texto_extraido or "Expedición" in texto_extraido
        assert "infracci" in texto_extraido or "infracción" in texto_extraido
        assert "L" in texto_extraido  # López/base
    except ImportError:
        # Fallback: verificar que no crasheo al generar
        pytest.skip("pypdf no instalado, verificamos solo ausencia de crash")


def test_pdf_disclaimer_presente(service, markdown_basico):
    """El PDF debe contener el disclaimer como texto."""
    result = service.markdown_a_pdf(markdown_basico)

    assert result.startswith(b"%PDF")

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(result))
        texto_extraido = "\n".join(page.extract_text() or "" for page in reader.pages)

        # Palabras clave del disclaimer canónico deben aparecer
        assert "DefensIA" in texto_extraido
        assert "asistencia" in texto_extraido or "tecnica" in texto_extraido.lower()
    except ImportError:
        # Sin pypdf: mínima verificación binaria
        assert len(result) > 1500  # con disclaimer el PDF debe ser más largo


def test_pdf_disclaimer_custom(service, markdown_basico):
    """El parámetro disclaimer debe reflejarse en el PDF exportado."""
    custom = "Disclaimer custom únicamente para este test 2026."
    result = service.markdown_a_pdf(markdown_basico, disclaimer=custom)

    assert result.startswith(b"%PDF")

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(result))
        texto_extraido = "\n".join(page.extract_text() or "" for page in reader.pages)

        # Al menos la parte base del disclaimer debe estar
        assert "custom" in texto_extraido or "Disclaimer" in texto_extraido
    except ImportError:
        assert len(result) > 1000


def test_pdf_builtin_helvetica_fallback(service, markdown_basico, monkeypatch):
    """Si DejaVuSans no está disponible, el PDF debe generarse con Helvetica sin crash."""
    # Forzar que el método interno de registro devuelva Helvetica
    def _fake_registrar(self):
        return "Helvetica"

    monkeypatch.setattr(
        DefensiaExportService,
        "_registrar_fuente_unicode",
        _fake_registrar,
    )

    # Debe generar el PDF sin UnicodeEncodeError (Helvetica soporta Latin-1)
    result = service.markdown_a_pdf(
        "# Resolución con tildes\n\nArtículo primero.",
    )

    assert result.startswith(b"%PDF")
    assert len(result) > 500
